#!/usr/bin/env python3
"""PLI 仕様書兼取扱説明書 Word (.docx) 生成スクリプト"""

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

# ── カラー定数 ──
NAVY = RGBColor(0x1a, 0x3a, 0x6a)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
DGRAY = RGBColor(0x66, 0x66, 0x66)
ACCENT = RGBColor(0x3a, 0x7a, 0xbf)

def set_cell_shading(cell, color_hex):
    """セルの背景色を設定"""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
    for edge, val in kwargs.items():
        element = parse_xml(
            f'<w:{edge} {nsdecls("w")} w:val="{val.get("val", "single")}" '
            f'w:sz="{val.get("sz", "4")}" w:space="0" '
            f'w:color="{val.get("color", "CCCCCC")}"/>'
        )
        tcBorders.append(element)
    tcPr.append(tcBorders)

def make_header_cell(cell, text):
    """テーブルヘッダーセル"""
    set_cell_shading(cell, "1a3a6a")
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.color.rgb = WHITE
    run.font.size = Pt(9)
    run.font.name = "Hiragino Kaku Gothic Pro"

def make_body_cell(cell, text, bold=False, center=False):
    """テーブルボディセル"""
    p = cell.paragraphs[0]
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(9)
    run.font.name = "Hiragino Mincho Pro"

def add_styled_table(doc, headers, rows, col_widths=None):
    """スタイル付きテーブル"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # ヘッダー
    for i, h in enumerate(headers):
        make_header_cell(table.cell(0, i), h)

    # ボディ
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.cell(r_idx + 1, c_idx)
            make_body_cell(cell, val)
            # 偶数行に薄い背景
            if r_idx % 2 == 1:
                set_cell_shading(cell, "f8f6f0")

    # 列幅
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)

    # 罫線
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '  <w:top w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
        '  <w:left w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
        '  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
        '  <w:right w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
        '  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
        '  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
        '</w:tblBorders>'
    )
    tblPr.append(borders)

    return table

def add_section_heading(doc, text):
    """紺色帯の見出し"""
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    set_cell_shading(cell, "1a3a6a")
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.color.rgb = WHITE
    run.font.size = Pt(13)
    run.font.name = "Hiragino Kaku Gothic Pro"
    doc.add_paragraph()  # spacer

def add_h2(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = NAVY
    run.font.name = "Hiragino Kaku Gothic Pro"
    p.space_before = Pt(12)

def add_h3(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = ACCENT
    run.font.name = "Hiragino Kaku Gothic Pro"
    p.space_before = Pt(8)

def add_body(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(9.5)
    run.font.name = "Hiragino Mincho Pro"
    p.paragraph_format.space_after = Pt(4)
    return p

def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.clear()
    run = p.add_run(text)
    run.font.size = Pt(9.5)
    run.font.name = "Hiragino Mincho Pro"
    return p

def add_code(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(8.5)
    run.font.name = "Courier New"
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    p.paragraph_format.left_indent = Cm(1)
    # 背景色
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="f0ede5" w:val="clear"/>')
    p.paragraph_format.element.get_or_add_pPr().append(shd)
    return p

def add_note(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(8.5)
    run.font.color.rgb = DGRAY
    run.font.name = "Hiragino Mincho Pro"
    p.paragraph_format.left_indent = Cm(1)
    return p


def main():
    doc = Document()

    # ── ページ設定 (A4) ──
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

    # ── デフォルトフォント ──
    style = doc.styles["Normal"]
    style.font.name = "Hiragino Mincho Pro"
    style.font.size = Pt(10)

    # ══════════════════════════════════════════════
    #  表紙
    # ══════════════════════════════════════════════
    for _ in range(4):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("PLI - Private Link Interpreter")
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = NAVY
    run.font.name = "Hiragino Kaku Gothic Pro"

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("秘匿通訳支援システム")
    run.font.size = Pt(12)
    run.font.color.rgb = DGRAY

    doc.add_paragraph()  # spacer

    # 水平線
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        '  <w:bottom w:val="single" w:sz="6" w:space="1" w:color="CCCCCC"/>'
        '</w:pBdr>'
    )
    pPr.append(pBdr)

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("仕様書 兼 取扱説明書")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = NAVY

    for _ in range(3):
        doc.add_paragraph()

    # 表紙の情報テーブル
    info = [
        ("バージョン", "2.0.0"),
        ("対応OS", "macOS 12+ (Apple Silicon)"),
        ("開発者", "関 智幸 (Tomoyuki Seki)"),
        ("作成日", "2026年3月14日"),
        ("Bundle ID", "com.seki.pli"),
    ]
    t = doc.add_table(rows=len(info), cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (k, v) in enumerate(info):
        c0, c1 = t.cell(i, 0), t.cell(i, 1)
        set_cell_shading(c0, "f0ede5")
        make_body_cell(c0, k, bold=True, center=True)
        make_body_cell(c1, v)
        c0.width = Cm(4)
        c1.width = Cm(8)
    # テーブル罫線
    tbl = t._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '  <w:top w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
        '  <w:left w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
        '  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
        '  <w:right w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
        '  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
        '  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
        '</w:tblBorders>'
    )
    tblPr.append(borders)

    doc.add_page_break()

    # ══════════════════════════════════════════════
    #  目次
    # ══════════════════════════════════════════════
    p = doc.add_paragraph()
    run = p.add_run("目次")
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = NAVY
    run.font.name = "Hiragino Kaku Gothic Pro"
    p.space_after = Pt(12)

    toc_items = [
        (False, "1.  概要"),
        (False, "2.  システム要件"),
        (False, "3.  インストールと起動"),
        (False, "4.  画面構成"),
        (True,  "  4.1  弁護人ウィンドウ（メインコンソール）"),
        (True,  "  4.2  被疑者ウィンドウ（表示用）"),
        (False, "5.  翻訳エンジン"),
        (True,  "  5.1  Hybrid（OPUS-MT + NLLB）"),
        (True,  "  5.2  NLLB単体"),
        (True,  "  5.3  LLM（llama.cpp）"),
        (False, "6.  音声認識（STT）"),
        (False, "7.  定型文テンプレート"),
        (False, "8.  固有名詞辞書（グロッサリー）"),
        (False, "9.  辞書検索"),
        (False, "10. セッション管理・エクスポート"),
        (False, "11. 秘匿モード（Hide / Panic）"),
        (False, "12. キーボードショートカット"),
        (False, "13. 設定項目一覧"),
        (False, "14. 対応言語一覧"),
        (False, "15. ファイル構成"),
        (False, "16. トラブルシューティング"),
    ]
    for is_sub, text in toc_items:
        p = doc.add_paragraph()
        run = p.add_run(text)
        if is_sub:
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
            p.paragraph_format.left_indent = Cm(1.2)
        else:
            run.font.size = Pt(10)
            run.font.color.rgb = NAVY
            run.bold = True
            p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.space_before = Pt(2)

    doc.add_page_break()

    # ══════════════════════════════════════════════
    #  1. 概要
    # ══════════════════════════════════════════════
    add_section_heading(doc, "1. 概要")
    add_body(doc,
        "PLI（Private Link Interpreter）は、弁護人と外国人被疑者の接見時に使用する"
        "リアルタイム通訳支援アプリケーションです。"
    )
    add_body(doc,
        "弁護人が日本語で話す／入力すると、自動的に相手方の言語に翻訳して被疑者画面に表示します。"
        "逆に被疑者が外国語で話した内容は日本語に翻訳して弁護人画面に表示します。"
    )
    add_h3(doc, "主な特徴")
    for f in [
        "Apple Silicon GPU 加速の音声認識（mlx-whisper）",
        "複数翻訳エンジン対応（OPUS-MT / NLLB / LLM）",
        "40以上の言語に対応",
        "固有名詞辞書によるカスタマイズ可能な翻訳",
        "法律分野に特化した定型文テンプレート",
        "秘匿モード（Hide）とパニックボタン",
        "会話ログの保存・エクスポート",
        "完全オフライン動作（モデルダウンロード後）",
    ]:
        add_bullet(doc, f)

    # ══════════════════════════════════════════════
    #  2. システム要件
    # ══════════════════════════════════════════════
    add_section_heading(doc, "2. システム要件")
    add_styled_table(doc, ["項目", "要件"], [
        ["OS", "macOS 12 (Monterey) 以上"],
        ["CPU", "Apple Silicon (M1/M2/M3/M4)"],
        ["メモリ", "8GB 以上（推奨16GB以上）"],
        ["ストレージ", "約3GB（アプリ＋翻訳モデル）"],
        ["Python", "3.10 以上（開発モード時）"],
        ["マイク", "内蔵マイクまたは外付けマイク"],
    ], [4, 13])
    add_note(doc, "※ LLMエンジン使用時は追加で4〜64GBのメモリが必要です。")

    # ══════════════════════════════════════════════
    #  3. インストールと起動
    # ══════════════════════════════════════════════
    add_section_heading(doc, "3. インストールと起動")

    add_h3(doc, "3.1 アプリ版（.app）で起動")
    add_body(doc, "dist/PLI.app をFinderでダブルクリックするか、ターミナルから以下を実行します。")
    add_code(doc, "open dist/PLI.app")
    add_note(doc, "※ 初回起動時に「開発元不明」警告が表示されます。右クリック →「開く」で許可してください。")

    add_h3(doc, "3.2 開発モードで起動")
    add_code(doc, "python main.py            # モックモード（テスト用）")
    add_code(doc, "python main.py --real      # 実モデル使用")

    add_h3(doc, "3.3 起動オプション")
    add_styled_table(doc, ["オプション", "説明"], [
        ["--real", "実モデルを使用（デフォルトはモック）"],
        ["--display MODE", "表示モード: auto / switch / unified / dual"],
        ["--engine ENGINE", "翻訳エンジン: auto / llm / nllb / hybrid"],
        ["--model PATH", "LLMモデルファイルのパスを指定"],
        ["--n-ctx N", "LLMコンテキスト長を指定"],
    ], [4, 13])

    add_h3(doc, "3.4 再ビルド（ソース変更後）")
    add_body(doc, "ソースコードを変更した場合、.appに反映するには再ビルドが必要です。")
    add_code(doc, "pyinstaller PLI.spec")
    add_note(doc, "※ python main.py で直接起動する場合はビルド不要です。")

    doc.add_page_break()

    # ══════════════════════════════════════════════
    #  4. 画面構成
    # ══════════════════════════════════════════════
    add_section_heading(doc, "4. 画面構成")

    add_h2(doc, "4.1 弁護人ウィンドウ（メインコンソール）")
    add_body(doc,
        "弁護人が操作するメインウィンドウです。会話ログの表示、テキスト入力、"
        "音声認識の制御、各種設定を行います。"
    )

    add_h3(doc, "画面レイアウト")
    add_styled_table(doc, ["領域", "内容"], [
        ["メニューバー", "言語(L) / 表示(V) / セッション(S) / 録音(R) / 音声認識(M) / テスト(T) / 設定(O) / ヘルプ(H)"],
        ["ツールバー", "モードインジケーター、読込状態、対象言語表示"],
        ["会話ログ", "吹き出し形式の会話履歴（弁護人=紺色、被疑者=緑色）"],
        ["承認パネル", "被疑者発言の確認ボックス（OK / やり直し / 編集ボタン）"],
        ["入力エリア", "定型文ボタン / 辞書ボタン / テキスト入力欄 / 送信ボタン"],
        ["ステータスバー", "セッション番号、録音サイズ、STTモード表示"],
    ], [3, 14])

    add_h3(doc, "表示モード")
    add_styled_table(doc, ["モード", "説明"], [
        ["switch", "フルスクリーン切替（弁護人画面 / 被疑者画面をF3で切替）"],
        ["unified", "左右分割（左に弁護人コンソール、右に被疑者パネル）"],
        ["split", "スプリッター分割（弁護人・被疑者を自由にリサイズ）"],
    ], [3, 14])

    add_h3(doc, "会話バブルの操作")
    add_styled_table(doc, ["操作", "内容"], [
        ["右クリック", "コピーメニュー（原文+訳文 / 原文のみ / 訳文のみ）"],
        ["修正ボタン", "弁護人発言の翻訳を手動修正"],
        ["取消ボタン", "発言の翻訳を取り消し"],
        ["編集ボタン", "被疑者発言の翻訳を手動修正"],
    ], [3, 14])

    add_h2(doc, "4.2 被疑者ウィンドウ（表示用）")
    add_body(doc,
        "被疑者に見せる表示専用ウィンドウです。タイプライター効果で翻訳結果を表示します。"
        "操作ボタンはなく、表示のみの画面です。"
    )
    add_styled_table(doc, ["要素", "説明"], [
        ["メッセージ表示", "色分けされた吹き出し（弁護人発言=紺、被疑者発言=緑）"],
        ["タイプライター効果", "弁護人の翻訳が一文字ずつ表示される演出"],
        ["ステータスバナー", "多言語対応の状態メッセージ（確認中/確定/修正中 等）"],
        ["自動スクロール", "新しいメッセージに自動追従"],
    ], [4, 13])

    doc.add_page_break()

    # ══════════════════════════════════════════════
    #  5. 翻訳エンジン
    # ══════════════════════════════════════════════
    add_section_heading(doc, "5. 翻訳エンジン")
    add_body(doc, "PLIは3種類の翻訳エンジンに対応しています。設定メニューから切り替え可能です。")

    add_h2(doc, "5.1 Hybrid（OPUS-MT + NLLB）【推奨】")
    add_body(doc,
        "OPUS-MTモデルを主翻訳エンジンとして使用し、"
        "対応していない言語ペアではNLLBにフォールバックします。最も高精度な翻訳が可能です。"
    )
    add_body(doc, "翻訳フロー: 日本語 → (OPUS-MT) → 英語 → (OPUS-MT) → 対象言語")
    add_note(doc, "※ 英語を中間言語（ピボット）として使用する2段階翻訳方式です。")

    add_h2(doc, "5.2 NLLB 単体")
    add_body(doc,
        "Meta社のNLLB-200モデルをCTranslate2で最適化して使用します。"
        "軽量でメモリ消費が少なく、CPU環境でも動作します。"
    )

    add_h2(doc, "5.3 LLM（llama.cpp）")
    add_body(doc,
        "大規模言語モデル（Qwen2.5等）を使用した翻訳です。"
        "構文チェック機能が利用可能ですが、大量のメモリを必要とします。"
    )
    add_styled_table(doc, ["モデルサイズ", "必要メモリ", "特徴"], [
        ["7B (Q4_K_M)", "約8GB", "基本的な翻訳"],
        ["14B (Q4_K_M)", "約12GB", "高品質な翻訳"],
        ["32B (Q4_K_M)", "約24GB", "専門分野対応"],
        ["72B (Q3_K_M)", "約40GB", "最高品質"],
    ], [4, 3, 10])

    # ══════════════════════════════════════════════
    #  6. 音声認識
    # ══════════════════════════════════════════════
    add_section_heading(doc, "6. 音声認識（STT）")
    add_body(doc,
        "Apple Silicon GPU加速のmlx-whisperを使用したリアルタイム音声認識です。"
        "マイクから取得した音声を自動でテキストに変換します。"
    )

    add_h3(doc, "操作方法")
    add_styled_table(doc, ["操作", "説明"], [
        ["Space / メニュー", "STTの開始・停止を切替"],
        ["Command+6", "自動言語検出モード"],
        ["Command+7", "弁護人入力モード（日本語固定）"],
        ["Command+8", "被疑者入力モード（外国語固定）"],
    ], [4, 13])

    add_h3(doc, "感度設定")
    add_styled_table(doc, ["レベル", "説明"], [
        ["high (1.3x)", "小さな声も検出（静かな環境向け）"],
        ["normal (1.8x)", "標準（推奨）"],
        ["low (2.5x)", "周囲の騒音が大きい環境向け"],
    ], [4, 13])

    add_h3(doc, "テンポ設定")
    add_styled_table(doc, ["レベル", "説明"], [
        ["slow", "無音判定 1.2秒 / 最小発話 0.5秒（ゆっくり話す人向け）"],
        ["normal", "無音判定 0.8秒 / 最小発話 0.3秒（標準）"],
        ["fast", "無音判定 0.5秒 / 最小発話 0.15秒（早口の人向け）"],
    ], [3, 14])
    add_note(doc, "※ STT動作中は入力欄の背景が黄色に変わります。")

    doc.add_page_break()

    # ══════════════════════════════════════════════
    #  7. 定型文テンプレート
    # ══════════════════════════════════════════════
    add_section_heading(doc, "7. 定型文テンプレート")
    add_body(doc,
        "入力エリアの「定型文」ボタンから、法律分野で頻出する定型文を選択して"
        "ワンクリックで翻訳・送信できます。"
    )

    add_h3(doc, "組込みカテゴリ")
    add_styled_table(doc, ["カテゴリ", "含まれるフレーズ"], [
        ["権利告知", "黙秘権の告知 / 弁護人選任権 / 供述の自由 / 接見交通権"],
        ["手続説明", "勾留の流れ / 保釈の説明 / 起訴・不起訴 / 裁判の流れ / 取調べの注意"],
        ["接見時の定型句", "挨拶 / 体調確認 / 次回面会 / 家族への連絡 / 終了の挨拶"],
    ], [4, 13])

    add_h3(doc, "カスタマイズ")
    add_body(doc, "設定メニュー →「定型文編集」から、カテゴリの追加・削除、フレーズの編集が可能です。")
    add_body(doc,
        "また、~/pli-models/定型文.docx をWordで直接編集することもできます。"
        "見出し1がカテゴリ名、2列の表（ラベル｜本文）がフレーズとして読み込まれます。"
    )
    add_note(doc, "読み込み優先順: ユーザーdocx → ユーザーjson → 同梱json")

    # ══════════════════════════════════════════════
    #  8. グロッサリー
    # ══════════════════════════════════════════════
    add_section_heading(doc, "8. 固有名詞辞書（グロッサリー）")
    add_body(doc,
        "人名や組織名などの固有名詞を正確に翻訳するための辞書機能です。"
        "翻訳エンジンに送る前に日本語の固有名詞を対応する外国語表記に直接置換し、"
        "翻訳後に正しく保持されているかチェックします。"
    )

    add_h3(doc, "動作の流れ")
    add_styled_table(doc, ["ステップ", "内容"], [
        ["1. 前処理", "入力テキスト中の固有名詞をローマ字/英語表記に置換"],
        ["2. 翻訳", "置換済みテキストを翻訳エンジンに送信"],
        ["3. 後処理", "翻訳結果に固有名詞が保持されているか確認。欠落時は強制置換"],
    ], [3, 14])

    add_h3(doc, "辞書の編集")
    add_body(doc,
        "設定メニュー →「固有名詞辞書」から、日本語と外国語のペアを追加・削除できます。"
        "type: \"name\" のエントリのみが固有名詞として処理されます。"
    )

    add_h3(doc, "登録例")
    add_styled_table(doc, ["日本語", "外国語"], [
        ["関智幸", "Tomoyuki Seki"],
        ["東京弁護士会", "Tokyo Bar Association"],
    ], [5, 12])

    # ══════════════════════════════════════════════
    #  9. 辞書検索
    # ══════════════════════════════════════════════
    add_section_heading(doc, "9. 辞書検索")
    add_body(doc,
        "入力エリアの「辞書」ボタンから、単語やフレーズの翻訳を検索できます。"
        "翻訳エンジンを使用して双方向（日本語→外国語、外国語→日本語）の"
        "リアルタイム検索が可能です。"
    )
    add_note(doc, "※ 検索結果は参考用です。会話ログには追加されません。")

    doc.add_page_break()

    # ══════════════════════════════════════════════
    #  10. セッション管理
    # ══════════════════════════════════════════════
    add_section_heading(doc, "10. セッション管理・エクスポート")

    add_h3(doc, "セッション")
    add_body(doc,
        "アプリ起動からの一連の会話を「セッション」として管理します。"
        "セッションメニューから新規セッションの開始や終了が可能です。"
    )

    add_h3(doc, "エクスポート形式")
    add_styled_table(doc, ["形式", "内容"], [
        ["JSON保存", "全データ（原文/訳文/中間英語/タイムスタンプ/ルート等）を構造化保存"],
        ["テキスト出力", "人間可読な形式（話者ラベル + 原文 → 訳文）で出力"],
    ], [3, 14])

    add_h3(doc, "録音モード")
    add_styled_table(doc, ["モード", "説明"], [
        ["OFF", "録音しない（デフォルト）"],
        ["VOLATILE", "RAM上に一時保存（アプリ終了で消滅）"],
        ["SAVE", "WAVファイルとして ~/pli-recordings/ に保存"],
    ], [3, 14])

    # ══════════════════════════════════════════════
    #  11. 秘匿モード
    # ══════════════════════════════════════════════
    add_section_heading(doc, "11. 秘匿モード（Hide / Panic）")
    add_body(doc, "接見中に第三者の目がある場合に、アプリの存在を隠す機能です。")

    add_h3(doc, "Hide モード（Command+1）")
    add_body(doc, "画面を即座にダミーPDF表示に切り替えます。トグル操作で元の画面に復帰できます。")
    add_body(doc, "オプション: 会話ログの消去 / 録音バッファの消去を設定可能")

    add_h3(doc, "Panic モード（Command+2）")
    add_body(doc, "不可逆のデータ破棄を行います。以下を即座に実行します。")
    for item in [
        "録音バッファをゼロで上書き消去",
        "会話ログを完全消去",
        "STTを停止",
        "画面をダミー表示に切替",
    ]:
        add_bullet(doc, item)
    add_note(doc, "Panicモードで消去されたデータは復元できません。")

    # ══════════════════════════════════════════════
    #  12. ショートカット
    # ══════════════════════════════════════════════
    add_section_heading(doc, "12. キーボードショートカット")
    add_styled_table(doc, ["ショートカット", "機能"], [
        ["Command+1", "秘匿モード切替（Hide）"],
        ["Command+2", "パニックボタン（データ全消去）"],
        ["Command+3", "被疑者画面の表示切替 / 埋込パネル切替"],
        ["Command+5", "音声認識 ON/OFF"],
        ["Command+6", "STT: 自動言語検出モード"],
        ["Command+7", "STT: 弁護人入力モード（日本語固定）"],
        ["Command+8", "STT: 被疑者入力モード（外国語固定）"],
        ["Command+D", "テスト: 被疑者発言シミュレーション"],
        ["Command+/", "ショートカット一覧を表示"],
        ["Space", "音声認識 ON/OFF（入力欄未選択時）"],
        ["Enter", "テキスト送信"],
    ], [4, 13])

    doc.add_page_break()

    # ══════════════════════════════════════════════
    #  13. 設定項目
    # ══════════════════════════════════════════════
    add_section_heading(doc, "13. 設定項目一覧")
    add_body(doc, "メニューバーの「設定(O)」から各種設定を変更できます。")
    add_styled_table(doc, ["項目", "説明"], [
        ["文字サイズ", "3段階のフォントスケーリング（小/中/大）"],
        ["LLMモデル", "llama.cppモデルファイルの選択"],
        ["コンテキスト長", "LLMのコンテキストトークン数"],
        ["翻訳エンジン", "Hybrid / NLLB / LLM の切替"],
        ["NLLBモデル", "NLLBモデルサイズの選択"],
        ["OPUS-MTモデル", "OPUS-MTモデルの管理"],
        ["定型文編集", "定型文テンプレートの追加・編集・削除"],
        ["固有名詞辞書", "グロッサリーの追加・編集・削除"],
        ["ダミーPDF", "秘匿モード用のPDFファイル選択"],
    ], [4, 13])

    add_h3(doc, "設定ファイルの保存先")
    add_code(doc, "~/pli-models/")
    add_styled_table(doc, ["ファイル", "内容"], [
        ["last_engine.txt", "最後に使用した翻訳エンジン"],
        ["last_model.txt", "最後に使用したLLMモデルパス"],
        ["last_n_ctx.txt", "最後に使用したコンテキスト長"],
        ["last_nllb_model.txt", "最後に使用したNLLBモデル"],
    ], [5, 12])

    # ══════════════════════════════════════════════
    #  14. 対応言語
    # ══════════════════════════════════════════════
    add_section_heading(doc, "14. 対応言語一覧")
    add_body(doc, "言語(L)メニューから対象言語を選択できます。品質ランクは以下の通りです。")
    add_styled_table(doc, ["ランク", "説明"], [
        ["(OPUS直通)", "日本語との直接OPUS-MTモデルあり。最高品質"],
        ["Tier A", "英語ピボット経由。OPUS-MTで高品質"],
        ["Tier B", "NLLBフォールバック。実用レベル"],
    ], [3, 14])

    add_h3(doc, "主な対応言語")
    add_styled_table(doc, ["言語", "Language", "ランク"], [
        ["英語", "English", "OPUS直通"],
        ["中国語(簡体)", "Chinese", "OPUS直通"],
        ["韓国語", "Korean", "OPUS直通"],
        ["ベトナム語", "Vietnamese", "Tier A"],
        ["ポルトガル語", "Portuguese", "Tier A"],
        ["スペイン語", "Spanish", "Tier A"],
        ["タガログ語", "Tagalog", "Tier A"],
        ["ネパール語", "Nepali", "Tier B"],
        ["ミャンマー語", "Burmese", "Tier B"],
        ["タイ語", "Thai", "Tier A"],
        ["インドネシア語", "Indonesian", "Tier A"],
        ["フランス語", "French", "OPUS直通"],
    ], [4, 4, 9])
    add_note(doc, "※ 上記以外にも40以上の言語に対応しています。")

    doc.add_page_break()

    # ══════════════════════════════════════════════
    #  15. ファイル構成
    # ══════════════════════════════════════════════
    add_section_heading(doc, "15. ファイル構成")
    add_styled_table(doc, ["ファイル", "説明"], [
        ["main.py", "アプリケーションエントリーポイント"],
        ["core/interpreter.py", "翻訳エンジン・グロッサリー処理"],
        ["core/stt_listener.py", "音声認識（mlx-whisper）"],
        ["core/recorder.py", "音声録音"],
        ["ui/attorney_window.py", "弁護人ウィンドウ UI"],
        ["ui/defendant_window.py", "被疑者ウィンドウ UI"],
        ["data/phrases.json", "定型文テンプレートデータ"],
        ["data/glossary.json", "固有名詞辞書データ"],
        ["assets/PLI.icns", "macOSアプリアイコン"],
        ["PLI.spec", "PyInstallerビルド設定"],
        ["pyproject.toml", "Pythonプロジェクト設定"],
        ["requirements.txt", "依存パッケージ一覧"],
        ["run.sh", "開発用起動スクリプト"],
    ], [5, 12])

    # ══════════════════════════════════════════════
    #  16. トラブルシューティング
    # ══════════════════════════════════════════════
    add_section_heading(doc, "16. トラブルシューティング")
    add_styled_table(doc, ["症状", "対処法"], [
        ["「開発元不明」で起動できない", "Finderで右クリック →「開く」で許可"],
        ["翻訳が遅い", "Hybridエンジンの初回起動時はモデルダウンロードが必要。2回目以降は高速"],
        ["マイクが認識されない", "システム設定 → プライバシーとセキュリティ → マイク で許可を確認"],
        ["固有名詞が正しく翻訳されない", "設定 → 固有名詞辞書に登録。type: name が設定されているか確認"],
        ["メモリ不足エラー", "翻訳エンジンをHybridまたはNLLBに変更（LLMより軽量）"],
        ["音声認識が途切れる", "音声認識メニューからテンポ設定を「slow」に変更"],
        [".appに変更が反映されない", "pyinstaller PLI.spec で再ビルドが必要"],
    ], [5, 12])

    # ── フッター ──
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("本書は PLI v2.0.0 の仕様に基づいて作成されています。")
    run.font.size = Pt(8.5)
    run.font.color.rgb = DGRAY
    p = doc.add_paragraph()
    run = p.add_run("作成: 2026年3月14日  |  PLI - Private Link Interpreter")
    run.font.size = Pt(8.5)
    run.font.color.rgb = DGRAY

    # ── 保存 ──
    output_path = os.path.join(os.path.dirname(__file__), "PLI_仕様書_取扱説明書.docx")
    doc.save(output_path)
    print(f"Word生成完了: {output_path}")
    return output_path


if __name__ == "__main__":
    main()
