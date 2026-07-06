"""
PLI TranslationPanel - 弁護人テキスト入力・翻訳制御ウィジェット
AttorneyWindow から抽出した入力・翻訳・定型文・承認パネルを独立 QWidget 化

Copyright (c) 2025-2026 中野通り法律事務所 弁護士 関智之（東京弁護士会所属）（東京弁護士会所属）(Tomoyuki Seki)
All rights reserved.
"""

import os
from typing import Optional

from PySide6.QtWidgets import (
    QWidget, QVBoxLayout, QHBoxLayout,
    QLabel, QPushButton, QLineEdit, QFrame, QMenu, QStatusBar,
)
from PySide6.QtCore import Qt, Signal, Slot
from PySide6.QtGui import QFont

from core.interpreter import Utterance, Speaker
from core.logging_setup import get_logger

logger = get_logger(__name__)

# ---------------------------------------------------------------------------
# スタイル定数 — 平成初期レトロ（暖灰・紺・深緑）
# ---------------------------------------------------------------------------

_BG       = "#d6d2c8"
_SURFACE  = "#e8e4da"
_SUNKEN   = "#c4c0b4"
_RAISED   = "#e0dcd2"
_RAISED_L = "#f2eee6"
_RAISED_D = "#9e9a8e"
_FIELD    = "#fffff4"
_TEXT     = "#1a1a10"
_DIM      = "#6a6658"
_ACCENT   = "#1a3a6a"
_DEF_CLR  = "#2a6a30"
_WARN     = "#8a3a1a"
_BTN_OK   = "#3a6a44"
_BTN_NG   = "#7a3a3a"
_BTN_EDIT = "#5a5a3a"
_BTN_SEND = "#3a4a6a"
_BTN_TEXT = "#f0ece0"


def _raised_border(bg=_SURFACE):
    return (
        f"background-color: {bg};"
        f"border-top: 1px solid {_RAISED_L};"
        f"border-left: 1px solid {_RAISED_L};"
        f"border-bottom: 1px solid {_RAISED_D};"
        f"border-right: 1px solid {_RAISED_D};"
    )


def _sunken_border(bg=_SUNKEN):
    return (
        f"background-color: {bg};"
        f"border-top: 1px solid {_RAISED_D};"
        f"border-left: 1px solid {_RAISED_D};"
        f"border-bottom: 1px solid {_RAISED_L};"
        f"border-right: 1px solid {_RAISED_L};"
    )


def _make_btn(label: str, bg: str) -> QPushButton:
    btn = QPushButton(label)
    btn.setCursor(Qt.PointingHandCursor)
    btn.setStyleSheet(f"""
        QPushButton {{
            background-color: {bg}; color: {_BTN_TEXT};
            border-top: 1px solid {_RAISED_L};
            border-left: 1px solid {_RAISED_L};
            border-bottom: 1px solid {_RAISED_D};
            border-right: 1px solid {_RAISED_D};
            padding: 4px 14px;
            font-size: 11px;
        }}
        QPushButton:pressed {{
            border-top: 1px solid {_RAISED_D};
            border-left: 1px solid {_RAISED_D};
            border-bottom: 1px solid {_RAISED_L};
            border-right: 1px solid {_RAISED_L};
        }}
    """)
    return btn


# ---------------------------------------------------------------------------
# TranslationPanel — 入力・翻訳・定型文・承認パネル
# ---------------------------------------------------------------------------

class TranslationPanel(QWidget):
    """弁護人テキスト入力と翻訳制御ウィジェット

    Signals:
        translation_requested(str): テキスト入力→翻訳要求
        ok_clicked(): 承認パネルOK
        retry_clicked(): 承認パネルやり直し
        edit_clicked(): 承認パネル手動修正
        phrase_send_requested(str): 定型文テキスト送信要求
    """

    translation_requested = Signal(str)
    ok_clicked = Signal()
    retry_clicked = Signal()
    edit_clicked = Signal()
    phrase_send_requested = Signal(str)

    def __init__(self, parent: Optional[QWidget] = None):
        super().__init__(parent)
        self._pending_utterance: Optional[Utterance] = None
        self._pending_is_attorney: bool = False
        self._translation_available: bool = False
        self._setup_ui()

    # ----- UI構築 -----

    def _setup_ui(self):
        layout = QVBoxLayout(self)
        layout.setSpacing(0)
        layout.setContentsMargins(0, 0, 0, 0)

        # === 承認パネル ===
        self.approval_frame = QFrame()
        self.approval_frame.setStyleSheet(
            f"{_raised_border(_SURFACE)} margin: 0 6px;"
        )
        self.approval_frame.setVisible(False)
        approval_layout = QVBoxLayout(self.approval_frame)
        approval_layout.setContentsMargins(10, 8, 10, 8)
        approval_layout.setSpacing(6)

        self.appr_header = QLabel("相手の発言 — 確認")
        self.appr_header.setStyleSheet(
            f"color: {_WARN}; font-size: 11px; font-weight: bold; border: none;"
        )
        self.appr_header.setFont(QFont("Menlo", 9))
        approval_layout.addWidget(self.appr_header)

        self.pending_label = QLabel("")
        self.pending_label.setWordWrap(True)
        self.pending_label.setTextInteractionFlags(
            Qt.TextSelectableByMouse | Qt.TextSelectableByKeyboard
        )
        self.pending_label.setCursor(Qt.IBeamCursor)
        self.pending_label.setStyleSheet(
            f"color: {_TEXT}; font-size: 12px; padding: 4px;"
            f"{_sunken_border(_FIELD)}"
        )
        approval_layout.addWidget(self.pending_label)

        btn_row = QHBoxLayout()
        btn_row.setSpacing(4)
        self.ok_btn = _make_btn("OK (O)", _BTN_OK)
        self.ok_btn.clicked.connect(self._on_ok)
        btn_row.addWidget(self.ok_btn)
        self.retry_btn = _make_btn("やり直し (R)", _BTN_NG)
        self.retry_btn.clicked.connect(self._on_retry)
        btn_row.addWidget(self.retry_btn)
        self.edit_btn = _make_btn("手動修正 (E)", _BTN_EDIT)
        self.edit_btn.clicked.connect(self._on_manual_edit)
        btn_row.addWidget(self.edit_btn)
        btn_row.addStretch()
        approval_layout.addLayout(btn_row)
        layout.addWidget(self.approval_frame)

        # === 入力エリア ===
        input_frame = QFrame()
        input_frame.setStyleSheet(
            f"{_raised_border(_SURFACE)} margin: 2px 6px 4px 6px;"
        )
        input_layout = QHBoxLayout(input_frame)
        input_layout.setContentsMargins(8, 6, 8, 6)
        input_layout.setSpacing(4)

        # 定型文ボタン
        self._phrase_btn = QPushButton("① 定型文")
        self._phrase_btn.setToolTip("定型文テンプレート")
        self._phrase_btn.setStyleSheet(f"""
            QPushButton {{
                background-color: {_ACCENT}; color: white;
                border-top: 1px solid {_RAISED_L};
                border-left: 1px solid {_RAISED_L};
                border-bottom: 1px solid {_RAISED_D};
                border-right: 1px solid {_RAISED_D};
                padding: 3px 8px; font-size: 11px; font-weight: bold;
            }}
            QPushButton:hover {{ background-color: #2a4a7a; }}
        """)
        self._phrase_btn.clicked.connect(self._on_phrase_menu)
        input_layout.addWidget(self._phrase_btn)

        # 入力フィールド
        self.input_field = QLineEdit()
        self.input_field.setPlaceholderText("日本語を入力してEnter...")
        self._input_normal_style = f"""
            QLineEdit {{
                background-color: {_FIELD}; color: {_TEXT};
                border-top: 1px solid {_RAISED_D};
                border-left: 1px solid {_RAISED_D};
                border-bottom: 1px solid {_RAISED_L};
                border-right: 1px solid {_RAISED_L};
                padding: 4px 8px; font-size: 13px;
                selection-background-color: {_ACCENT};
                selection-color: white;
            }}
        """
        self._input_stt_style = f"""
            QLineEdit {{
                background-color: #fffde0; color: {_TEXT};
                border: 2px solid #c8a830;
                padding: 3px 7px; font-size: 13px;
                selection-background-color: {_ACCENT};
                selection-color: white;
            }}
        """
        self.input_field.setStyleSheet(self._input_normal_style)
        self.input_field.returnPressed.connect(self._on_send_attorney)
        input_layout.addWidget(self.input_field)

        # 送信ボタン
        send_btn = _make_btn("送信", _BTN_SEND)
        send_btn.clicked.connect(self._on_send_attorney)
        input_layout.addWidget(send_btn)

        layout.addWidget(input_frame)

    # ----- 公開スロット (AttorneyWindow から呼ばれる) -----

    @Slot(object)
    def set_translation_result(self, utt: Utterance):
        """翻訳結果を承認パネルに表示する"""
        self._pending_utterance = utt
        self._pending_is_attorney = (utt.speaker == Speaker.ATTORNEY)

        header = "弁護人の発言 — 確認" if self._pending_is_attorney else "相手の発言 — 確認"
        self.appr_header.setText(header)

        display = utt.translated or utt.original
        self.pending_label.setText(display)
        self.approval_frame.setVisible(True)

    @Slot(bool)
    def set_loading(self, loading: bool):
        """翻訳中の入力フィールド無効化"""
        self.input_field.setEnabled(not loading)
        self._phrase_btn.setEnabled(not loading)
        if loading:
            self.input_field.setPlaceholderText("翻訳中...")
        else:
            self.input_field.setPlaceholderText("日本語を入力してEnter...")

    def set_translation_available(self, available: bool):
        """翻訳エンジンの利用可否を設定"""
        self._translation_available = available

    def set_input_text(self, text: str):
        """入力フィールドにテキストをセット（手動修正用）"""
        self.input_field.setText(text)
        self.input_field.setFocus()
        self.input_field.selectAll()

    def set_stt_style(self, active: bool):
        """STT入力中のスタイル切替"""
        if active:
            self.input_field.setStyleSheet(self._input_stt_style)
        else:
            self.input_field.setStyleSheet(self._input_normal_style)

    def clear_approval(self):
        """承認パネルを非表示にする"""
        self._pending_utterance = None
        self._pending_is_attorney = False
        self.approval_frame.setVisible(False)

    @property
    def pending_utterance(self) -> Optional[Utterance]:
        return self._pending_utterance

    @property
    def pending_is_attorney(self) -> bool:
        return self._pending_is_attorney

    # ----- 内部イベントハンドラ -----

    def _on_send_attorney(self):
        text = self.input_field.text().strip()
        if not text:
            return
        if not self._translation_available:
            return
        self.input_field.clear()
        self.input_field.setStyleSheet(self._input_normal_style)
        self.translation_requested.emit(text)

    def _on_ok(self):
        if self._pending_utterance:
            self.ok_clicked.emit()

    def _on_retry(self):
        self.retry_clicked.emit()

    def _on_manual_edit(self):
        if not self._pending_utterance:
            return
        self.edit_clicked.emit()

    # ----- 定型文 -----

    def _on_phrase_menu(self):
        """定型文カテゴリメニューをポップアップ"""
        data = self._load_phrases()
        menu = QMenu(self)
        menu.setStyleSheet(f"""
            QMenu {{
                background-color: {_SURFACE}; color: {_TEXT};
                font-size: 12px; border: 1px solid {_RAISED_D}; padding: 2px 0;
            }}
            QMenu::item {{ padding: 4px 20px 4px 10px; }}
            QMenu::item:selected {{ background-color: {_ACCENT}; color: white; }}
            QMenu::separator {{ height: 1px; background: {_RAISED_D}; margin: 2px 4px; }}
        """)
        for cat in data.get("categories", []):
            icon = cat.get("icon", "")
            sub = menu.addMenu(f'{icon} {cat["name"]}')
            sub.setStyleSheet(menu.styleSheet())
            for ph in cat.get("phrases", []):
                action = sub.addAction(ph["label"])
                action.triggered.connect(
                    lambda checked, t=ph["text"]: self._send_phrase(t)
                )
        menu.exec(self._phrase_btn.mapToGlobal(self._phrase_btn.rect().bottomLeft()))

    def _send_phrase(self, text: str):
        """定型文を入力欄にセットして送信"""
        self.input_field.setText(text)
        self._on_send_attorney()

    def _load_phrases(self) -> dict:
        """定型文を読み込む（.docx優先 -> .json フォールバック）"""
        user_docx = os.path.expanduser("~/pli-models/定型文.docx")
        user_json = os.path.expanduser("~/pli-models/phrases.json")
        bundled_json = os.path.join(
            os.path.dirname(os.path.dirname(__file__)), "data", "phrases.json"
        )
        if os.path.exists(user_docx):
            try:
                return self._load_phrases_docx(user_docx)
            except Exception as e:
                logger.warning("docx読込失敗, JSONにフォールバック: %s", e)
        import json
        path = user_json if os.path.exists(user_json) else bundled_json
        try:
            with open(path, "r", encoding="utf-8") as f:
                return json.load(f)
        except Exception:
            return {"categories": []}

    @staticmethod
    def _load_phrases_docx(path: str) -> dict:
        """Wordファイルから定型文を読み込む

        フォーマット:
          見出し1 = カテゴリ名
          表 = 2列（ラベル | 本文）
        """
        from docx import Document
        doc = Document(path)
        categories = []
        current_cat = None

        for element in doc.element.body:
            tag = element.tag.split("}")[-1]

            if tag == "p":
                from docx.oxml.ns import qn
                pStyle = element.find(qn("w:pPr"))
                style_name = ""
                if pStyle is not None:
                    ps = pStyle.find(qn("w:pStyle"))
                    if ps is not None:
                        style_name = ps.get(qn("w:val"), "")

                if style_name in ("Heading1", "1", "heading 1", "見出し1", "見出し 1"):
                    text = element.text.strip() if hasattr(element, 'text') and element.text else ""
                    if not text:
                        texts = [r.text for r in element.findall(qn("w:r") + "/" + qn("w:t")) if r.text]
                        text = "".join(texts).strip()
                    if text:
                        icon, name = "", text
                        if len(text) > 1 and not text[0].isalnum() and text[0] not in "あ-ん":
                            for i, ch in enumerate(text):
                                if ch == " " or ch == "\u3000":
                                    icon = text[:i]
                                    name = text[i:].strip()
                                    break
                        current_cat = {"name": name, "icon": icon, "phrases": []}
                        categories.append(current_cat)

            elif tag == "tbl" and current_cat is not None:
                from docx.oxml.ns import qn
                rows = element.findall(qn("w:tr"))
                for row in rows:
                    cells = row.findall(qn("w:tc"))
                    if len(cells) >= 2:
                        label_texts = []
                        body_texts = []
                        for p in cells[0].findall(qn("w:p")):
                            for r in p.findall(qn("w:r")):
                                t = r.find(qn("w:t"))
                                if t is not None and t.text:
                                    label_texts.append(t.text)
                        for p in cells[1].findall(qn("w:p")):
                            for r in p.findall(qn("w:r")):
                                t = r.find(qn("w:t"))
                                if t is not None and t.text:
                                    body_texts.append(t.text)
                        label = "".join(label_texts).strip()
                        body = "".join(body_texts).strip()
                        if label and body and label != "ラベル" and label != "名前":
                            current_cat["phrases"].append(
                                {"label": label, "text": body}
                            )
        return {"categories": categories}

    def _save_phrases(self, data: dict):
        """定型文データを Word (.docx) で保存"""
        from docx import Document
        from docx.shared import Pt, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        user_path = os.path.expanduser("~/pli-models/定型文.docx")
        os.makedirs(os.path.dirname(user_path), exist_ok=True)

        doc = Document()

        title = doc.add_heading("PLI 定型文テンプレート", level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        note = doc.add_paragraph(
            "このファイルを直接編集して保存すると、PLIの定型文メニューに反映されます。\n"
            "カテゴリ＝「見出し1」、フレーズ＝表（ラベル｜本文）の形式です。"
        )
        note.style = doc.styles["Normal"]
        for run in note.runs:
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

        for cat in data.get("categories", []):
            icon = cat.get("icon", "")
            name = cat.get("name", "")
            heading_text = f"{icon} {name}".strip() if icon else name
            doc.add_heading(heading_text, level=1)

            phrases = cat.get("phrases", [])
            if not phrases:
                doc.add_paragraph("（フレーズなし）")
                continue

            table = doc.add_table(rows=1 + len(phrases), cols=2)
            table.style = "Table Grid"

            hdr = table.rows[0]
            hdr.cells[0].text = "ラベル"
            hdr.cells[1].text = "本文"
            for cell in hdr.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
                        run.font.size = Pt(10)

            for i, ph in enumerate(phrases):
                row = table.rows[i + 1]
                row.cells[0].text = ph.get("label", "")
                row.cells[1].text = ph.get("text", "")
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(10)

            for row in table.rows:
                row.cells[0].width = Cm(4)
                row.cells[1].width = Cm(12)

            doc.add_paragraph("")

        doc.save(user_path)
        return user_path
