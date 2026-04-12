"""
PLI Attorney Console - 弁護人コンソール画面
平成初期レトロUI — 一太郎的な業務用ソフトの佇まい

Copyright (c) 2025-2026 中野通り法律事務所 弁護士 関智之 (Tomoyuki Seki)
All rights reserved.

表示モード:
  split  — QSplitter左右分割（デュアルモニタープレビュー向け）
  switch — QStackedWidget全画面切替（F3で弁護人⇔被疑者トグル）
"""

import os
import queue
import threading
from dataclasses import dataclass

from PySide6.QtWidgets import (
    QMainWindow, QWidget, QVBoxLayout, QHBoxLayout,
    QLabel, QPushButton, QTextEdit, QLineEdit,
    QRadioButton, QButtonGroup, QFrame, QScrollArea,
    QSplitter, QStackedWidget, QStatusBar, QDialog, QTableWidget,
    QTableWidgetItem, QHeaderView, QAbstractItemView,
    QMenuBar, QMenu, QInputDialog, QFileDialog, QMessageBox,
    QListWidget, QListWidgetItem, QComboBox,
)
from PySide6.QtCore import Qt, Signal, Slot, QTimer
from PySide6.QtGui import QFont, QColor, QKeySequence, QShortcut, QAction, QActionGroup

from core.interpreter import Interpreter, Utterance, Speaker, SyntaxChunk, EngineType, SUPPORTED_LANGUAGES, get_language_name
from core.recorder import Recorder, RecordMode
from ui.defendant_window import DefendantPanel

from ui.conversation_bubble import ConversationBubble  # noqa: F401
from ui.dialogs import SyntaxCheckDialog, PhraseEditorDialog, GlossaryEditorDialog, DictionaryDialog  # noqa: F401

@dataclass(frozen=True)
class TranslationJob:
    job_id: int
    session_token: int
    kind: str
    text: str


# ---------------------------------------------------------------------------
# スタイル定数 — 平成初期レトロ（暖灰・紺・深緑）
# ---------------------------------------------------------------------------

_BG       = "#d6d2c8"
_SURFACE  = "#e8e4da"
_SUNKEN   = "#c4c0b4"
_RAISED   = "#e0dcd2"
_RAISED_L = "#f2eee6"
_RAISED_D = "#9e9a8e"
_FIELD    = "#fffff4"
_TEXT     = "#1a1a10"
_DIM      = "#6a6658"
_ACCENT   = "#1a3a6a"
_DEF_CLR  = "#2a6a30"
_WARN     = "#8a3a1a"
_BTN_OK   = "#3a6a44"
_BTN_NG   = "#7a3a3a"
_BTN_EDIT = "#5a5a3a"
_BTN_SEND = "#3a4a6a"
_BTN_TEXT = "#f0ece0"


def _raised_border(bg=_SURFACE):
    return (
        f"background-color: {bg};"
        f"border-top: 1px solid {_RAISED_L};"
        f"border-left: 1px solid {_RAISED_L};"
        f"border-bottom: 1px solid {_RAISED_D};"
        f"border-right: 1px solid {_RAISED_D};"
    )

def _sunken_border(bg=_SUNKEN):
    return (
        f"background-color: {bg};"
        f"border-top: 1px solid {_RAISED_D};"
        f"border-left: 1px solid {_RAISED_D};"
        f"border-bottom: 1px solid {_RAISED_L};"
        f"border-right: 1px solid {_RAISED_L};"
    )


def _make_btn(label: str, bg: str) -> QPushButton:
    btn = QPushButton(label)
    btn.setCursor(Qt.PointingHandCursor)
    btn.setStyleSheet(f"""
        QPushButton {{
            background-color: {bg}; color: {_BTN_TEXT};
            border-top: 1px solid {_RAISED_L};
            border-left: 1px solid {_RAISED_L};
            border-bottom: 1px solid {_RAISED_D};
            border-right: 1px solid {_RAISED_D};
            padding: 4px 14px;
            font-size: 11px;
        }}
        QPushButton:pressed {{
            border-top: 1px solid {_RAISED_D};
            border-left: 1px solid {_RAISED_D};
            border-bottom: 1px solid {_RAISED_L};
            border-right: 1px solid {_RAISED_L};
        }}
    """)
    return btn


from ui.font_config import fs as _fs
import ui.font_config as _font_cfg


# ---------------------------------------------------------------------------
# メインウィンドウ: 弁護人コンソール
# ---------------------------------------------------------------------------

class AttorneyWindow(QMainWindow):
    """弁護人コンソール メインウィンドウ"""

    send_to_defendant = Signal(str, str)
    stream_to_defendant = Signal(str)
    finish_defendant_stream = Signal()
    defendant_correction = Signal()
    defendant_retry = Signal()
    defendant_update_last = Signal(str)
    defendant_lang_change = Signal(str)
    request_hide = Signal()
    request_reveal = Signal()
    request_panic = Signal()
    clear_defendant = Signal()
    embedded_panel_toggled = Signal(bool)   # 埋め込みパネルON/OFF通知
    stt_result = Signal(str, str)           # STT結果 (text, lang)
    stt_toggled = Signal(bool)              # STT ON/OFF通知
    stt_sensitivity_changed = Signal(str)   # 感度プリセット (high/normal/low)
    stt_tempo_changed = Signal(str)         # テンポプリセット (slow/normal/fast)
    llm_model_changed = Signal(str)         # LLMモデル変更 (ファイルパス)
    llm_ctx_changed = Signal(int)            # コンテキスト長変更
    engine_type_changed = Signal(str)        # エンジン種別変更 (llm/nllb/hybrid)
    nllb_model_changed = Signal(str)         # NLLBモデル変更 (モデルキー)
    opus_download_requested = Signal(str)    # OPUS-MTペアダウンロード要求
    _defendant_translated = Signal(object)   # 被疑者翻訳完了（バックグラウンド→メイン）
    _attorney_translated = Signal(object)    # 弁護人翻訳完了（バックグラウンド→メイン）
    _manual_attorney_translated = Signal(int, object)
    _manual_attorney_failed = Signal(int, str)
    _attorney_translation_failed = Signal(str, str)
    _interpreter_stream_token_ready = Signal(object, str)
    _interpreter_utterance_ready = Signal(object)

    def __init__(self, interpreter: Interpreter, recorder: Recorder,
                 view_style: str = "split"):
        super().__init__()
        self.interpreter = interpreter
        self.recorder = recorder
        self._pending_utterance: Utterance | None = None
        self._pending_is_attorney: bool = False  # 承認パネルが弁護人用かどうか
        self._current_attorney_bubble: ConversationBubble | None = None
        self._last_attorney_bubble: ConversationBubble | None = None  # 修正/取消ボタン管理
        self._last_defendant_bubble: ConversationBubble | None = None
        self._editing_def_bubble: ConversationBubble | None = None
        self._session_count = 1
        self._dummy_pdf_path = ""
        self._pending_attorney_bubbles: dict[int, ConversationBubble] = {}
        self._translation_job_seq = 0
        self._translation_session_token = 0
        self._translation_state_lock = threading.Lock()
        self._cancelled_translation_job_ids: set[int] = set()
        self._translation_queue: queue.Queue[TranslationJob | None] = queue.Queue()
        self._embed_visible = False
        self._view_style = view_style      # "split" or "switch"

        self.setWindowTitle("PLI — Private Link Interpreter")
        self.setMinimumSize(700, 600)

        # 埋め込み用被疑者パネル
        self._defendant_panel = DefendantPanel()

        self._setup_ui()
        self._setup_shortcuts()
        self._setup_callbacks()
        self._connect_embedded_panel()

        # 翻訳完了シグナル（バックグラウンドスレッド→メインスレッド）
        self._defendant_translated.connect(self._on_defendant_translated)
        self._attorney_translated.connect(self._on_attorney_translated)
        self._manual_attorney_translated.connect(self._on_manual_attorney_translated)
        self._manual_attorney_failed.connect(self._on_manual_attorney_failed)
        self._attorney_translation_failed.connect(self._on_attorney_translation_failed)
        self._interpreter_stream_token_ready.connect(self._on_interpreter_stream_token)
        self._interpreter_utterance_ready.connect(self._on_interpreter_utterance)
        self._start_translation_worker()

    @property
    def defendant_panel(self) -> DefendantPanel:
        """埋め込み用被疑者パネルへのアクセス"""
        return self._defendant_panel

    def _connect_embedded_panel(self):
        """弁護人のシグナルを埋め込みパネルにも接続"""
        self.send_to_defendant.connect(self._defendant_panel.on_message)
        self.stream_to_defendant.connect(self._defendant_panel.on_stream_token)
        self.finish_defendant_stream.connect(self._defendant_panel.finish_stream)
        self.defendant_correction.connect(self._defendant_panel.on_correction)
        self.defendant_retry.connect(self._defendant_panel.on_retry)
        self.defendant_update_last.connect(self._defendant_panel.on_update_last)
        self.defendant_lang_change.connect(self._defendant_panel.on_language_change)
        self.clear_defendant.connect(self._defendant_panel.on_clear)

    def _setup_ui(self):
        self.setStyleSheet(f"background-color: {_BG};")

        central = QWidget()
        self.setCentralWidget(central)
        outer_layout = QVBoxLayout(central)
        outer_layout.setSpacing(0)
        outer_layout.setContentsMargins(0, 0, 0, 0)

        self._session_label = None

        # ===== ステルスモード用トップレベルスタック =====
        self._stealth_stack = QStackedWidget()
        self._is_stealth = False

        # ===== レイアウトコンテナ =====
        self._splitter = None

        # --- 実コンテンツラッパー ---
        real_content = QWidget()
        real_content.setStyleSheet(f"background-color: {_BG};")
        real_layout = QVBoxLayout(real_content)
        real_layout.setSpacing(0)
        real_layout.setContentsMargins(0, 0, 0, 0)

        # --- 左側: 弁護人コンテンツ ---
        attorney_widget = QWidget()
        attorney_widget.setStyleSheet(f"background-color: {_BG};")
        main_layout = QVBoxLayout(attorney_widget)
        main_layout.setSpacing(0)
        main_layout.setContentsMargins(0, 0, 0, 0)

        # ツールバー
        toolbar = QFrame()
        toolbar.setFixedHeight(28)
        toolbar.setStyleSheet(
            f"{_raised_border(_SURFACE)} padding: 2px 8px;"
        )
        tb_layout = QHBoxLayout(toolbar)
        tb_layout.setContentsMargins(8, 0, 8, 0)
        tb_layout.setSpacing(4)

        self._mode_label = QLabel("MOCK")
        self._mode_label.setFont(QFont("Menlo", 9))
        self._mode_label.setStyleSheet(
            f"color: {_DIM}; font-size: 10px; border: none;"
        )
        tb_layout.addWidget(self._mode_label)

        self._loading_label = QLabel("")
        self._loading_label.setFont(QFont("Menlo", 9))
        self._loading_label.setStyleSheet(
            f"color: {_WARN}; font-size: 10px; font-weight: bold; border: none;"
        )
        tb_layout.addWidget(self._loading_label)

        tb_layout.addStretch()

        self._lang_label = QLabel("English")
        self._lang_label.setFont(QFont("Menlo", 9))
        self._lang_label.setStyleSheet(
            f"color: {_ACCENT}; font-size: 10px; font-weight: bold; border: none;"
        )
        tb_layout.addWidget(self._lang_label)

        main_layout.addWidget(toolbar)

        # 会話ログ
        log_frame = QFrame()
        log_frame.setStyleSheet(
            f"{_sunken_border(_FIELD)} margin: 4px 6px;"
        )
        log_inner = QVBoxLayout(log_frame)
        log_inner.setContentsMargins(0, 0, 0, 0)

        # ログ上部ツールバー（全コピーボタン）
        log_toolbar = QHBoxLayout()
        log_toolbar.setContentsMargins(8, 2, 8, 0)
        log_toolbar.addStretch()
        copy_all_btn = QPushButton("📋 全会話コピー")
        copy_all_btn.setToolTip("会話ログ全体をクリップボードにコピー")
        copy_all_btn.setStyleSheet(f"""
            QPushButton {{
                background-color: {_SURFACE}; color: {_DIM};
                font-size: 11px; padding: 2px 10px;
                border: 1px solid {_RAISED_D}; border-radius: 3px;
            }}
            QPushButton:hover {{ background-color: {_RAISED_L}; color: {_TEXT}; }}
        """)
        copy_all_btn.clicked.connect(self._copy_all_conversation)
        log_toolbar.addWidget(copy_all_btn)
        log_inner.addLayout(log_toolbar)

        self.scroll_area = QScrollArea()
        self.scroll_area.setWidgetResizable(True)
        self.scroll_area.setStyleSheet(f"""
            QScrollArea {{ border: none; background-color: {_FIELD}; }}
            QScrollBar:vertical {{
                background: {_SURFACE}; width: 14px;
                border-left: 1px solid {_RAISED_D};
            }}
            QScrollBar::handle:vertical {{
                {_raised_border(_BG)} min-height: 20px;
            }}
            QScrollBar::add-line:vertical, QScrollBar::sub-line:vertical {{
                height: 14px; {_raised_border(_BG)}
            }}
        """)
        self.log_container = QWidget()
        self.log_container.setStyleSheet(f"background-color: {_FIELD};")
        self.log_layout = QVBoxLayout(self.log_container)
        self.log_layout.setAlignment(Qt.AlignTop)
        self.log_layout.setSpacing(0)
        self.log_layout.setContentsMargins(8, 4, 8, 4)
        self.scroll_area.setWidget(self.log_container)
        log_inner.addWidget(self.scroll_area)
        main_layout.addWidget(log_frame, stretch=1)

        # 承認パネル
        self.approval_frame = QFrame()
        self.approval_frame.setStyleSheet(
            f"{_raised_border(_SURFACE)} margin: 0 6px;"
        )
        self.approval_frame.setVisible(False)
        approval_layout = QVBoxLayout(self.approval_frame)
        approval_layout.setContentsMargins(10, 8, 10, 8)
        approval_layout.setSpacing(6)

        self.appr_header = QLabel("相手の発言 — 確認")
        self.appr_header.setStyleSheet(
            f"color: {_WARN}; font-size: 11px; font-weight: bold; border: none;"
        )
        self.appr_header.setFont(QFont("Menlo", 9))
        approval_layout.addWidget(self.appr_header)

        self.pending_label = QLabel("")
        self.pending_label.setWordWrap(True)
        self.pending_label.setTextInteractionFlags(
            Qt.TextSelectableByMouse | Qt.TextSelectableByKeyboard
        )
        self.pending_label.setCursor(Qt.IBeamCursor)
        self.pending_label.setStyleSheet(
            f"color: {_TEXT}; font-size: 12px; padding: 4px;"
            f"{_sunken_border(_FIELD)}"
        )
        approval_layout.addWidget(self.pending_label)

        btn_row = QHBoxLayout()
        btn_row.setSpacing(4)
        self.ok_btn = _make_btn("OK (O)", _BTN_OK)
        self.ok_btn.clicked.connect(self._on_ok)
        btn_row.addWidget(self.ok_btn)
        self.retry_btn = _make_btn("やり直し (R)", _BTN_NG)
        self.retry_btn.clicked.connect(self._on_retry)
        btn_row.addWidget(self.retry_btn)
        self.edit_btn = _make_btn("手動修正 (E)", _BTN_EDIT)
        self.edit_btn.clicked.connect(self._on_manual_edit)
        btn_row.addWidget(self.edit_btn)
        btn_row.addStretch()
        approval_layout.addLayout(btn_row)
        main_layout.addWidget(self.approval_frame)

        # 入力エリア
        input_frame = QFrame()
        input_frame.setStyleSheet(
            f"{_raised_border(_SURFACE)} margin: 2px 6px 4px 6px;"
        )
        input_layout = QHBoxLayout(input_frame)
        input_layout.setContentsMargins(8, 6, 8, 6)
        input_layout.setSpacing(4)

        # 定型文ボタン（①紺色）
        self._phrase_btn = QPushButton("① 定型文")
        self._phrase_btn.setToolTip("定型文テンプレート")
        self._phrase_btn.setStyleSheet(f"""
            QPushButton {{
                background-color: {_ACCENT}; color: white;
                border-top: 1px solid {_RAISED_L};
                border-left: 1px solid {_RAISED_L};
                border-bottom: 1px solid {_RAISED_D};
                border-right: 1px solid {_RAISED_D};
                padding: 3px 8px; font-size: 11px; font-weight: bold;
            }}
            QPushButton:hover {{ background-color: #2a4a7a; }}
        """)
        self._phrase_btn.clicked.connect(self._on_phrase_menu)
        input_layout.addWidget(self._phrase_btn)

        # 辞書ボタン（②緑色）
        self._dict_btn = QPushButton("② 辞書")
        self._dict_btn.setToolTip("辞書検索")
        self._dict_btn.setStyleSheet(f"""
            QPushButton {{
                background-color: {_BTN_OK}; color: white;
                border-top: 1px solid {_RAISED_L};
                border-left: 1px solid {_RAISED_L};
                border-bottom: 1px solid {_RAISED_D};
                border-right: 1px solid {_RAISED_D};
                padding: 3px 8px; font-size: 11px; font-weight: bold;
            }}
            QPushButton:hover {{ background-color: #4a7a54; }}
        """)
        self._dict_btn.clicked.connect(self._on_dict_dialog)
        input_layout.addWidget(self._dict_btn)

        self.input_field = QLineEdit()
        self.input_field.setPlaceholderText("日本語を入力してEnter...")
        self._input_normal_style = f"""
            QLineEdit {{
                background-color: {_FIELD}; color: {_TEXT};
                border-top: 1px solid {_RAISED_D};
                border-left: 1px solid {_RAISED_D};
                border-bottom: 1px solid {_RAISED_L};
                border-right: 1px solid {_RAISED_L};
                padding: 4px 8px; font-size: 13px;
                selection-background-color: {_ACCENT};
                selection-color: white;
            }}
        """
        self._input_stt_style = f"""
            QLineEdit {{
                background-color: #fffde0; color: {_TEXT};
                border: 2px solid #c8a830;
                padding: 3px 7px; font-size: 13px;
                selection-background-color: {_ACCENT};
                selection-color: white;
            }}
        """
        self.input_field.setStyleSheet(self._input_normal_style)
        self.input_field.returnPressed.connect(self._on_send_attorney)
        input_layout.addWidget(self.input_field)

        send_btn = _make_btn("送信", _BTN_SEND)
        send_btn.clicked.connect(self._on_send_attorney)
        input_layout.addWidget(send_btn)
        main_layout.addWidget(input_frame)

        # --- レイアウトモードに応じてコンテナ構築 ---
        self._attorney_widget = attorney_widget  # 参照保持

        # 両モード共通: QSplitter で管理
        self._splitter = QSplitter(Qt.Horizontal)
        self._splitter.setStyleSheet(
            f"QSplitter::handle {{ background: {_RAISED_D}; width: 3px; }}"
        )
        self._splitter.addWidget(attorney_widget)
        self._splitter.addWidget(self._defendant_panel)

        if self._view_style == "switch":
            # ----- switch: ⌘3で 弁護人→被疑者→分割 を循環 -----
            self._switch_state = 0   # 0=弁護人, 1=被疑者, 2=分割
            self._defendant_panel.setVisible(False)
            self._splitter.setStretchFactor(0, 1)
            self._splitter.setStretchFactor(1, 1)
        else:
            # ----- split: 埋め込みトグル -----
            self._defendant_panel.setVisible(False)
            self._splitter.setStretchFactor(0, 3)
            self._splitter.setStretchFactor(1, 2)

        real_layout.addWidget(self._splitter)

        # ----- ステルスモード: ダミー画面 -----
        stealth_widget = self._create_stealth_widget()
        self._stealth_stack.addWidget(real_content)     # page 0 = 実画面
        self._stealth_stack.addWidget(stealth_widget)   # page 1 = ダミー
        self._stealth_stack.setCurrentIndex(0)
        outer_layout.addWidget(self._stealth_stack)

        # ----- メニューバー -----
        self._setup_menubar()

        # ----- ステータスバー -----
        self.status_bar = QStatusBar()
        self.status_bar.setStyleSheet(f"""
            QStatusBar {{
                background-color: {_BG}; color: {_DIM};
                font-size: 10px;
                border-top: 1px solid {_RAISED_D};
                padding: 1px 4px;
            }}
            QStatusBar::item {{ border: none; }}
        """)
        self.setStatusBar(self.status_bar)
        self.status_bar.showMessage("待機中")

        # STTモード常設表示
        self._stt_mode_label = QLabel("")
        self._stt_mode_label.setFont(QFont("Menlo", 9))
        self._stt_mode_label.setStyleSheet(
            f"color: {_DIM}; font-size: 10px; padding: 0 6px;"
            f"border-left: 1px solid {_RAISED_D};"
        )
        self._stt_mode_label.setVisible(False)
        self.status_bar.addPermanentWidget(self._stt_mode_label)

        self._session_label = QLabel(f"Session #{self._session_count:02d}")
        self._session_label.setStyleSheet(
            f"color: {_DIM}; font-size: 10px; padding: 0 6px;"
            f"border-left: 1px solid {_RAISED_D};"
            f"border-right: 1px solid {_RAISED_L};"
        )
        self._session_label.setFont(QFont("Menlo", 9))
        self.status_bar.addPermanentWidget(self._session_label)

        self.rec_size_label = QLabel("")
        self.rec_size_label.setStyleSheet(
            f"color: {_WARN}; font-size: 10px; padding: 0 6px;"
            f"border-left: 1px solid {_RAISED_D};"
        )
        self.rec_size_label.setFont(QFont("Menlo", 9))
        self.status_bar.addPermanentWidget(self.rec_size_label)

        self._rec_timer = QTimer()
        self._rec_timer.timeout.connect(self._update_rec_size)
        self._rec_timer.start(1000)

    def _setup_menubar(self):
        menubar = self.menuBar()
        menubar.setStyleSheet(f"""
            QMenuBar {{
                background-color: {_SURFACE}; color: {_TEXT};
                font-size: 12px;
                border-bottom: 1px solid {_RAISED_D};
                padding: 1px 2px;
            }}
            QMenuBar::item {{ padding: 3px 8px; }}
            QMenuBar::item:selected {{
                background: {_ACCENT}; color: white;
            }}
            QMenu {{
                background-color: {_SURFACE}; color: {_TEXT};
                font-size: 12px;
                border: 1px solid {_RAISED_D};
                padding: 2px 0;
            }}
            QMenu::item {{ padding: 4px 20px 4px 10px; }}
            QMenu::item:selected {{
                background-color: {_ACCENT}; color: white;
            }}
            QMenu::separator {{
                height: 1px; background: {_RAISED_D}; margin: 2px 4px;
            }}
        """)

        # ----- 言語(L) -----
        lang_menu = menubar.addMenu("言語(&L)")
        self._lang_action_group = QActionGroup(self)
        self._lang_action_group.setExclusive(True)

        # OPUS-MT直接ペアがある言語を判定
        try:
            from core.opus_downloader import OPUS_MODELS
            _opus_direct = set()
            for key in OPUS_MODELS:
                parts = key.split("-")
                if len(parts) == 2:
                    _opus_direct.update(parts)
            _opus_direct.discard("en")
            _opus_direct.discard("ja")
        except Exception:
            _opus_direct = set()

        tier_order = {"◎": 0, "○": 1, "△": 2}
        sorted_langs = sorted(
            SUPPORTED_LANGUAGES.items(),
            key=lambda x: (tier_order.get(x[1]["tier"], 9), x[1]["name"]),
        )
        current_tier = None
        for code, info in sorted_langs:
            tier = info["tier"]
            if current_tier is not None and tier != current_tier:
                lang_menu.addSeparator()
            current_tier = tier
            # 英語経由の言語のみタグ表示
            if code != "en" and code not in _opus_direct:
                route_tag = " [英語経由]"
            else:
                route_tag = ""
            label = f'{info["tier"]} {info["name"]}（{info["native"]}）{route_tag}'
            action = QAction(label, self, checkable=True)
            if code == "en":
                action.setChecked(True)
            action.triggered.connect(lambda checked, c=code: self._on_language_change(c))
            self._lang_action_group.addAction(action)
            lang_menu.addAction(action)

        # ----- 表示(V) -----
        view_menu = menubar.addMenu("表示(&V)")

        if self._view_style == "switch":
            self._embed_action = QAction("相手画面に切替  ⌘3", self)
            self._embed_action.triggered.connect(self._toggle_embedded_panel)
            view_menu.addAction(self._embed_action)
        else:
            self._embed_action = QAction("相手画面を埋め込み  ⌘3", self, checkable=True)
            self._embed_action.triggered.connect(self._toggle_embedded_panel)
            view_menu.addAction(self._embed_action)

        # ----- セッション(S) -----
        session_menu = menubar.addMenu("セッション(&S)")

        save_json_action = QAction("💾 記録を保存 (JSON)(&J)", self)
        save_json_action.triggered.connect(self._on_save_json)
        session_menu.addAction(save_json_action)

        save_text_action = QAction("📝 記録をエクスポート (テキスト)(&T)", self)
        save_text_action.triggered.connect(self._on_save_text)
        session_menu.addAction(save_text_action)

        session_menu.addSeparator()

        end_action = QAction("セッション終了(&E)", self)
        end_action.triggered.connect(self._on_end_session)
        session_menu.addAction(end_action)
        session_menu.addSeparator()

        hide_action = QAction("画面を隠す  ⌘1", self)
        hide_action.triggered.connect(self._on_f1)
        session_menu.addAction(hide_action)

        panic_action = QAction("緊急消去    ⌘2", self)
        panic_action.triggered.connect(self._on_f2)
        session_menu.addAction(panic_action)

        # ----- 録音(R) -----
        rec_menu = menubar.addMenu("録音(&R)")
        self._rec_action_group = QActionGroup(self)
        self._rec_action_group.setExclusive(True)

        rec_off = QAction("OFF", self, checkable=True)
        rec_off.setChecked(True)
        rec_off.triggered.connect(lambda: self._on_rec_mode(RecordMode.OFF))
        self._rec_action_group.addAction(rec_off)
        rec_menu.addAction(rec_off)

        rec_volatile = QAction("一時録音（揮発）", self, checkable=True)
        rec_volatile.triggered.connect(lambda: self._on_rec_mode(RecordMode.VOLATILE))
        self._rec_action_group.addAction(rec_volatile)
        rec_menu.addAction(rec_volatile)

        rec_save = QAction("録音保存", self, checkable=True)
        rec_save.triggered.connect(lambda: self._on_rec_mode(RecordMode.SAVE))
        self._rec_action_group.addAction(rec_save)
        rec_menu.addAction(rec_save)

        # ----- 音声認識(M) -----
        stt_menu = menubar.addMenu("音声認識(&M)")
        self._stt_action = QAction("🎤 マイクON  Space", self, checkable=True)
        self._stt_action.triggered.connect(self._toggle_stt)
        stt_menu.addAction(self._stt_action)
        self._stt_active = False

        # 言語モード: AUTO / 弁護人(JA強制) / 相手(外国語強制)
        stt_menu.addSeparator()
        self._stt_lang_mode = "auto"  # "auto" / "attorney" / "defendant"
        self._stt_lang_group = QActionGroup(self)
        _lang_auto = QAction("自動判定 (AUTO)  ⌘6", self, checkable=True)
        _lang_auto.setChecked(True)
        _lang_auto.triggered.connect(lambda: self._set_stt_lang_mode("auto"))
        _lang_atty = QAction("弁護人として入力  ⌘7", self, checkable=True)
        _lang_atty.triggered.connect(lambda: self._set_stt_lang_mode("attorney"))
        _lang_def = QAction("相手の発言として入力  ⌘8", self, checkable=True)
        _lang_def.triggered.connect(lambda: self._set_stt_lang_mode("defendant"))
        for a in (_lang_auto, _lang_atty, _lang_def):
            self._stt_lang_group.addAction(a)
            stt_menu.addAction(a)

        # マイク感度
        stt_menu.addSeparator()
        sens_menu = stt_menu.addMenu("🎚 マイク感度")
        self._sens_group = QActionGroup(self)
        self._sens_presets = {
            "high":   "高感度（小声でも拾う）",
            "normal": "標準",
            "low":    "低感度（ノイズ環境）",
        }
        for key, label in self._sens_presets.items():
            act = QAction(label, self, checkable=True)
            if key == "normal":
                act.setChecked(True)
            act.triggered.connect(lambda checked, k=key: self._set_stt_sensitivity(k))
            self._sens_group.addAction(act)
            sens_menu.addAction(act)

        # 発話テンポ
        tempo_menu = stt_menu.addMenu("🗣 発話テンポ")
        self._tempo_group = QActionGroup(self)
        self._tempo_presets = {
            "slow":   "ゆっくり（無音長め判定）",
            "normal": "標準",
            "fast":   "早口（短い間も区切らない）",
        }
        for key, label in self._tempo_presets.items():
            act = QAction(label, self, checkable=True)
            if key == "normal":
                act.setChecked(True)
            act.triggered.connect(lambda checked, k=key: self._set_stt_tempo(k))
            self._tempo_group.addAction(act)
            tempo_menu.addAction(act)

        # ----- テスト(T) -----
        test_menu = menubar.addMenu("テスト(&T)")
        def_sim_action = QAction("相手の発言を入力...  ⌘D", self)
        def_sim_action.triggered.connect(self._on_simulate_defendant)
        test_menu.addAction(def_sim_action)

        # ----- ヘルプ(H) -----
        help_menu = menubar.addMenu("ヘルプ(&H)")
        shortcut_help = QAction("ショートカット一覧  ⌘/", self)
        shortcut_help.triggered.connect(self._show_shortcut_help)
        help_menu.addAction(shortcut_help)
        help_menu.addSeparator()
        about_action = QAction("PLI について...", self)
        about_action.triggered.connect(self._show_about)
        help_menu.addAction(about_action)

        # ----- 設定(O) -----
        settings_menu = menubar.addMenu("設定(&O)")

        # フォントサイズ
        font_menu = settings_menu.addMenu("🔤 フォントサイズ")
        font_group = QActionGroup(self)
        font_group.setExclusive(True)
        for label, scale in [("小 (×0.7)", 0.7), ("やや小 (×0.85)", 0.85),
                              ("標準 (×1.0)", 1.0), ("やや大 (×1.2)", 1.2),
                              ("大 (×1.5)", 1.5), ("特大 (×2.0)", 2.0)]:
            act = QAction(label, self, checkable=True)
            if abs(scale - _font_cfg.font_scale) < 0.01:
                act.setChecked(True)
            act.triggered.connect(lambda checked, s=scale: self._set_font_scale(s))
            font_group.addAction(act)
            font_menu.addAction(act)
        settings_menu.addSeparator()

        # LLMモデル選択
        self._model_menu = settings_menu.addMenu("🤖 LLMモデル")
        self._model_group = QActionGroup(self)
        self._model_group.setExclusive(True)
        self._scan_and_populate_models()

        # コンテキスト長
        ctx_menu = settings_menu.addMenu("📏 コンテキスト長")
        self._ctx_group = QActionGroup(self)
        self._ctx_group.setExclusive(True)
        current_ctx = getattr(self.interpreter, '_n_ctx', 2048)
        self._ctx_presets = [
            (2048,  "2K  （省メモリ）"),
            (4096,  "4K  （標準）"),
            (8192,  "8K  （長い会話向け）"),
            (16384, "16K （大容量メモリ向け）"),
        ]
        for val, label in self._ctx_presets:
            act = QAction(label, self, checkable=True)
            if val == current_ctx:
                act.setChecked(True)
            act.triggered.connect(lambda checked, v=val: self._on_ctx_select(v))
            self._ctx_group.addAction(act)
            ctx_menu.addAction(act)

        settings_menu.addSeparator()

        # 翻訳エンジン選択
        engine_menu = settings_menu.addMenu("🔧 翻訳エンジン")
        self._engine_group = QActionGroup(self)
        self._engine_group.setExclusive(True)

        from core.interpreter import EngineType
        current_engine = getattr(self.interpreter, '_engine_type', EngineType.MOCK)

        self._engine_llm_action = QAction("LLM (llama.cpp) — GPU加速", self, checkable=True)
        if current_engine in (EngineType.LLM, EngineType.MOCK):
            self._engine_llm_action.setChecked(True)
        self._engine_llm_action.triggered.connect(lambda: self._on_engine_select("llm"))
        self._engine_group.addAction(self._engine_llm_action)
        engine_menu.addAction(self._engine_llm_action)

        self._engine_nllb_action = QAction("NLLB (軽量) — 8GB対応", self, checkable=True)
        if current_engine == EngineType.NLLB:
            self._engine_nllb_action.setChecked(True)
        self._engine_nllb_action.triggered.connect(lambda: self._on_engine_select("nllb"))
        self._engine_group.addAction(self._engine_nllb_action)
        engine_menu.addAction(self._engine_nllb_action)

        self._engine_hybrid_action = QAction("⚡ ハイブリッド (最高精度) — 32GB以上推奨", self, checkable=True)
        if current_engine == EngineType.HYBRID:
            self._engine_hybrid_action.setChecked(True)
        self._engine_hybrid_action.triggered.connect(lambda: self._on_engine_select("hybrid"))
        self._engine_group.addAction(self._engine_hybrid_action)
        engine_menu.addAction(self._engine_hybrid_action)

        # NLLBモデル選択
        self._nllb_menu = settings_menu.addMenu("🌐 NLLBモデル")
        self._nllb_group = QActionGroup(self)
        self._nllb_group.setExclusive(True)
        self._populate_nllb_models()

        # OPUS-MTペア管理
        self._opus_menu = settings_menu.addMenu("🔤 OPUS-MT言語ペア")
        self._populate_opus_models()

        settings_menu.addSeparator()

        self._hide_wipe_log = QAction("隠す時にログも消去", self, checkable=True)
        self._hide_wipe_log.setChecked(True)
        settings_menu.addAction(self._hide_wipe_log)
        self._hide_wipe_rec = QAction("隠す時に録音も消去", self, checkable=True)
        self._hide_wipe_rec.setChecked(True)
        settings_menu.addAction(self._hide_wipe_rec)
        settings_menu.addSeparator()
        edit_phrases_action = QAction("📋 定型文を編集", self)
        edit_phrases_action.triggered.connect(self._on_edit_phrases)
        settings_menu.addAction(edit_phrases_action)

        edit_glossary_action = QAction("📖 固有名詞辞書を編集", self)
        edit_glossary_action.triggered.connect(self._on_edit_glossary)
        settings_menu.addAction(edit_glossary_action)
        settings_menu.addSeparator()
        dummy_action = QAction("ダミーPDFを選択...", self)
        dummy_action.triggered.connect(self._on_select_dummy_pdf)
        settings_menu.addAction(dummy_action)

    def _on_select_dummy_pdf(self):
        from PySide6.QtWidgets import QFileDialog
        path, _ = QFileDialog.getOpenFileName(
            self, "ダミー表示用PDFを選択", "", "PDF Files (*.pdf)"
        )
        if path:
            self._dummy_pdf_path = path
            self.status_bar.showMessage(f"ダミーPDFを設定: {os.path.basename(path)}", 3000)

    def get_hide_preferences(self) -> dict[str, object]:
        return {
            "wipe_log_on_hide": getattr(self, "_hide_wipe_log", None).isChecked()
            if hasattr(self, "_hide_wipe_log") else True,
            "wipe_recording_on_hide": getattr(self, "_hide_wipe_rec", None).isChecked()
            if hasattr(self, "_hide_wipe_rec") else True,
            "dummy_pdf_path": self._dummy_pdf_path,
        }

    # ---- LLMモデル選択 ----

    # モデル名 → 人間向けラベルのマッピング
    _MODEL_LABELS = {
        # 72B — 64GB RAM向け
        "Qwen2.5-72B-Instruct-Q5_K_M": "⭐ Qwen2.5 72B Q5_K_M（最高品質・RAM 52GB〜）",
        "Qwen2.5-72B-Instruct-Q4_K_M": "⭐ Qwen2.5 72B Q4_K_M（高品質・RAM 42GB〜）",
        "Qwen2.5-72B-Instruct-Q3_K_M": "Qwen2.5 72B Q3_K_M（RAM 32GB〜）",
        # 32B — 32GB RAM向け
        "Qwen2.5-32B-Instruct-Q8_0":   "Qwen2.5 32B Q8_0  （ほぼ無劣化・RAM 34GB〜）",
        "Qwen2.5-32B-Instruct-Q6_K":   "Qwen2.5 32B Q6_K  （高品質・RAM 24GB〜）",
        "Qwen2.5-32B-Instruct-Q5_K_M": "Qwen2.5 32B Q5_K_M（推奨・RAM 21GB〜）",
        "Qwen2.5-32B-Instruct-Q4_K_M": "Qwen2.5 32B Q4_K_M（軽量・RAM 18GB〜）",
        "Qwen2.5-32B-Instruct-Q3_K_M": "Qwen2.5 32B Q3_K_M（最軽量・RAM 15GB〜）",
        # 14B — 16GB RAM向け
        "Qwen2.5-14B-Instruct-Q6_K":   "Qwen2.5 14B Q6_K  （16GB向け・高品質）",
        "Qwen2.5-14B-Instruct-Q5_K_M": "Qwen2.5 14B Q5_K_M（16GB向け・推奨）",
        "Qwen2.5-14B-Instruct-Q4_K_M": "Qwen2.5 14B Q4_K_M（16GB向け・軽量）",
        # 7B — 8GB RAM向け
        "Qwen2.5-7B-Instruct-Q6_K":    "Qwen2.5 7B  Q6_K  （8GB向け・高速）",
        "Qwen2.5-7B-Instruct-Q5_K_M":  "Qwen2.5 7B  Q5_K_M（8GB向け）",
    }

    def _scan_and_populate_models(self):
        """~/pli-models/ をスキャンしてモデルメニューを構築"""
        import os, glob
        models_dir = os.path.expanduser("~/pli-models")
        gguf_files = sorted(glob.glob(os.path.join(models_dir, "*.gguf")))

        # 現在ロード中のモデルパス
        current = getattr(self.interpreter, '_model_path', '') or ''

        if not gguf_files:
            no_model = QAction("（モデルが見つかりません）", self)
            no_model.setEnabled(False)
            self._model_menu.addAction(no_model)
            return

        for path in gguf_files:
            stem = os.path.splitext(os.path.basename(path))[0]
            label = self._MODEL_LABELS.get(stem, stem)
            act = QAction(label, self, checkable=True)
            if path == current or (current and os.path.basename(current) == os.path.basename(path)):
                act.setChecked(True)
            act.triggered.connect(lambda checked, p=path, s=stem: self._on_model_select(p, s))
            self._model_group.addAction(act)
            self._model_menu.addAction(act)

        # 区切り + 再スキャン
        self._model_menu.addSeparator()
        rescan = QAction("🔄 モデルを再スキャン", self)
        rescan.triggered.connect(self._rescan_models)
        self._model_menu.addAction(rescan)

    def _on_model_select(self, path: str, stem: str):
        """モデル選択時"""
        from PySide6.QtWidgets import QMessageBox
        reply = QMessageBox.question(
            self, "モデル切替",
            f"LLMモデルを切り替えます。\n\n"
            f"  {self._MODEL_LABELS.get(stem, stem)}\n\n"
            f"アプリを再起動して反映します。よろしいですか？",
            QMessageBox.Yes | QMessageBox.No, QMessageBox.Yes)
        if reply == QMessageBox.Yes:
            self.status_bar.showMessage(f"🤖 モデル変更: {stem}  → 再起動します...", 5000)
            self.llm_model_changed.emit(path)
        else:
            # キャンセル → チェックを元に戻す
            self._rescan_models()

    def _rescan_models(self):
        """モデルメニューを再構築"""
        self._model_menu.clear()
        self._model_group = QActionGroup(self)
        self._model_group.setExclusive(True)
        self._scan_and_populate_models()
        self.status_bar.showMessage("🔄 モデル一覧を更新しました", 2000)

    def _on_ctx_select(self, n_ctx: int):
        """コンテキスト長選択時"""
        from PySide6.QtWidgets import QMessageBox
        reply = QMessageBox.question(
            self, "コンテキスト長変更",
            f"コンテキスト長を {n_ctx:,} トークンに変更します。\n"
            f"アプリを再起動して反映します。よろしいですか？",
            QMessageBox.Yes | QMessageBox.No, QMessageBox.Yes)
        if reply == QMessageBox.Yes:
            self.status_bar.showMessage(f"📏 コンテキスト長: {n_ctx:,}  → 再起動します...", 5000)
            self.llm_ctx_changed.emit(n_ctx)
        else:
            # キャンセル → チェックを元に戻す
            current_ctx = getattr(self.interpreter, '_n_ctx', 2048)
            for act in self._ctx_group.actions():
                for val, label in self._ctx_presets:
                    if label == act.text() and val == current_ctx:
                        act.setChecked(True)

    # ---- 翻訳エンジン選択 ----

    def _on_engine_select(self, engine_type: str):
        """翻訳エンジン選択時"""
        from PySide6.QtWidgets import QMessageBox
        labels = {"llm": "LLM (llama.cpp)", "nllb": "NLLB (軽量)", "hybrid": "ハイブリッド (最高精度)"}
        label = labels.get(engine_type, engine_type)
        reply = QMessageBox.question(
            self, "翻訳エンジン変更",
            f"翻訳エンジンを {label} に変更します。\n"
            f"アプリを再起動して反映します。よろしいですか？",
            QMessageBox.Yes | QMessageBox.No, QMessageBox.Yes)
        if reply == QMessageBox.Yes:
            self.status_bar.showMessage(f"🔧 エンジン変更: {label}  → 再起動します...", 5000)
            self.engine_type_changed.emit(engine_type)
        else:
            # キャンセル → チェックを元に戻す
            from core.interpreter import EngineType
            current = getattr(self.interpreter, '_engine_type', EngineType.MOCK)
            if current == EngineType.NLLB:
                self._engine_nllb_action.setChecked(True)
            elif current == EngineType.HYBRID:
                self._engine_hybrid_action.setChecked(True)
            else:
                self._engine_llm_action.setChecked(True)

    def _populate_nllb_models(self):
        """NLLBモデルメニューを構築"""
        self._nllb_menu.clear()
        try:
            from core.nllb_downloader import NLLB_MODELS, is_downloaded, check_dependencies
            deps_ok, _ = check_dependencies()
        except ImportError:
            act = QAction("(nllb_downloader 読込エラー)", self)
            act.setEnabled(False)
            self._nllb_menu.addAction(act)
            return

        current_nllb = getattr(self.interpreter, '_nllb_model_dir', '')

        for key, info in NLLB_MODELS.items():
            downloaded = is_downloaded(key)
            label = info["label"]
            if downloaded:
                label = f"✅ {label}"
            else:
                label = f"⬇️ {label}  ({info['size_gb']:.1f}GB DL)"

            act = QAction(label, self, checkable=True)
            if downloaded and current_nllb and key in current_nllb:
                act.setChecked(True)
            if not deps_ok and not downloaded:
                act.setEnabled(False)
            act.triggered.connect(lambda checked, k=key, dl=downloaded: self._on_nllb_select(k, dl))
            self._nllb_group.addAction(act)
            self._nllb_menu.addAction(act)

        if not deps_ok:
            self._nllb_menu.addSeparator()
            dep_info = QAction("⚠️ pip install ctranslate2 transformers sentencepiece", self)
            dep_info.setEnabled(False)
            self._nllb_menu.addAction(dep_info)

    def _on_nllb_select(self, model_key: str, already_downloaded: bool):
        """NLLBモデル選択時"""
        from PySide6.QtWidgets import QMessageBox, QProgressDialog
        from core.nllb_downloader import NLLB_MODELS, is_downloaded, download_model

        if not already_downloaded:
            info = NLLB_MODELS[model_key]
            reply = QMessageBox.question(
                self, "NLLBモデルダウンロード",
                f"{info['label']}\n"
                f"サイズ: {info['size_gb']:.1f}GB\n"
                f"ダウンロードしますか？（時間がかかります）",
                QMessageBox.Yes | QMessageBox.No, QMessageBox.No)
            if reply != QMessageBox.Yes:
                self._populate_nllb_models()
                return

            # プログレスダイアログ
            progress = QProgressDialog("NLLBモデルをダウンロード中...", "キャンセル", 0, 100, self)
            progress.setWindowTitle("ダウンロード")
            progress.setMinimumDuration(0)
            progress.setValue(0)

            import threading
            def _do_download():
                try:
                    def on_progress(ratio):
                        progress.setValue(int(ratio * 100))
                    download_model(model_key, on_progress=on_progress)
                    progress.setValue(100)
                except Exception as e:
                    print(f"[error] NLLBダウンロード失敗: {e}")
                finally:
                    self._populate_nllb_models()

            threading.Thread(target=_do_download, daemon=True).start()
            progress.exec()

            if not is_downloaded(model_key):
                return

        # ダウンロード済み → エンジン切替
        self.status_bar.showMessage(f"🌐 NLLBモデル: {model_key}  → 再起動します...", 5000)
        self.nllb_model_changed.emit(model_key)

    def update_engine_mode(self, is_nllb: bool = False, is_hybrid: bool = False):
        """NLLB/ハイブリッドモード時にUIを調整（手動修正ボタンのグレーアウト等）"""
        no_syntax = is_nllb or is_hybrid
        self.edit_btn.setEnabled(not no_syntax)
        if is_nllb:
            self.edit_btn.setToolTip("NLLBモードでは構文チェックを利用できません")
        elif is_hybrid:
            self.edit_btn.setToolTip("ハイブリッドモードでは構文チェックを利用できません")
        else:
            self.edit_btn.setToolTip("")

    def _populate_opus_models(self):
        """OPUS-MT言語ペアメニューを構築"""
        self._opus_menu.clear()
        try:
            from core.opus_downloader import OPUS_MODELS, is_downloaded, check_dependencies
            deps_ok, _ = check_dependencies()
        except ImportError:
            act = QAction("(opus_downloader 読込エラー)", self)
            act.setEnabled(False)
            self._opus_menu.addAction(act)
            return

        # 優先度順にソート
        sorted_keys = sorted(OPUS_MODELS.keys(),
                             key=lambda k: (OPUS_MODELS[k].get("priority", 99), k))

        current_priority = None
        for key in sorted_keys:
            info = OPUS_MODELS[key]
            # 優先度グループごとにセパレータ
            if current_priority is not None and info.get("priority", 99) != current_priority:
                self._opus_menu.addSeparator()
            current_priority = info.get("priority", 99)

            downloaded = is_downloaded(key)
            label = info["label"]
            if downloaded:
                label = f"✅ {label}"
            else:
                label = f"⬇️ {label}  (~{info['ram_gb']:.1f}GB)"

            act = QAction(label, self)
            if not downloaded and deps_ok:
                act.triggered.connect(
                    lambda checked, k=key: self._on_opus_download(k))
            elif not deps_ok:
                act.setEnabled(False)
            else:
                act.setEnabled(True)
                act.setCheckable(True)
                act.setChecked(True)
            self._opus_menu.addAction(act)

        # --- マルチリンガルモデル（少数言語対応） ---
        self._opus_menu.addSeparator()
        mul_header = QAction("🌍 マルチリンガル（少数言語対応）", self)
        mul_header.setEnabled(False)
        self._opus_menu.addAction(mul_header)

        try:
            from core.opus_downloader import OPUS_MULTILINGUAL
            for mkey, minfo in OPUS_MULTILINGUAL.items():
                m_downloaded = is_downloaded(mkey)
                m_label = minfo["label"]
                if m_downloaded:
                    m_label = f"✅ {m_label}"
                else:
                    m_label = f"⬇️ {m_label}  (~{minfo['ram_gb']:.1f}GB)"
                m_act = QAction(m_label, self)
                if not m_downloaded and deps_ok:
                    m_act.triggered.connect(
                        lambda checked, k=mkey: self._on_opus_download(k))
                elif not deps_ok:
                    m_act.setEnabled(False)
                else:
                    m_act.setEnabled(True)
                    m_act.setCheckable(True)
                    m_act.setChecked(True)
                self._opus_menu.addAction(m_act)
        except ImportError:
            pass

        # 一括ダウンロード
        self._opus_menu.addSeparator()
        dl_all = QAction("📥 対象言語の全ペアをDL...", self)
        dl_all.triggered.connect(self._on_opus_download_all)
        if not deps_ok:
            dl_all.setEnabled(False)
        self._opus_menu.addAction(dl_all)

        # 統計情報
        downloaded_count = sum(1 for k in OPUS_MODELS if is_downloaded(k))
        try:
            from core.opus_downloader import OPUS_MULTILINGUAL
            mul_count = sum(1 for k in OPUS_MULTILINGUAL if is_downloaded(k))
            stat_text = f"({downloaded_count}/{len(OPUS_MODELS)} ペア + {mul_count}/{len(OPUS_MULTILINGUAL)} マルチリンガル DL済み)"
        except ImportError:
            stat_text = f"({downloaded_count}/{len(OPUS_MODELS)} ペアDL済み)"
        stat = QAction(stat_text, self)
        stat.setEnabled(False)
        self._opus_menu.addAction(stat)

        if not deps_ok:
            self._opus_menu.addSeparator()
            dep_info = QAction("⚠️ pip install ctranslate2 transformers sentencepiece huggingface_hub", self)
            dep_info.setEnabled(False)
            self._opus_menu.addAction(dep_info)

    def _on_opus_download(self, pair_key: str):
        """個別OPUS-MTペア or マルチリンガルモデルのダウンロード"""
        from PySide6.QtWidgets import QMessageBox, QProgressDialog
        from core.opus_downloader import OPUS_MODELS, OPUS_MULTILINGUAL, download_model, is_downloaded

        # ペア専用 or マルチリンガル
        if pair_key in OPUS_MODELS:
            info = OPUS_MODELS[pair_key]
        elif pair_key in OPUS_MULTILINGUAL:
            info = OPUS_MULTILINGUAL[pair_key]
        else:
            return
        reply = QMessageBox.question(
            self, "OPUS-MTモデルダウンロード",
            f"{info['label']}\n"
            f"RAM: ~{info['ram_gb']:.1f}GB\n"
            f"ダウンロード＆変換しますか？",
            QMessageBox.Yes | QMessageBox.No, QMessageBox.No)
        if reply != QMessageBox.Yes:
            return

        progress = QProgressDialog(f"OPUS-MT [{pair_key}] ダウンロード中...", "キャンセル", 0, 100, self)
        progress.setWindowTitle("ダウンロード")
        progress.setMinimumDuration(0)
        progress.setValue(0)

        import threading
        def _do_download():
            try:
                def on_progress(ratio):
                    progress.setValue(int(ratio * 100))
                download_model(pair_key, on_progress=on_progress)
                progress.setValue(100)
            except Exception as e:
                print(f"[error] OPUS-MT ダウンロード失敗: {e}")
            finally:
                self._populate_opus_models()

        threading.Thread(target=_do_download, daemon=True).start()
        progress.exec()

        if is_downloaded(pair_key):
            self.status_bar.showMessage(f"✅ OPUS-MT [{pair_key}] ダウンロード完了", 5000)

    def _on_opus_download_all(self):
        """対象言語の関連ペアを一括ダウンロード"""
        from PySide6.QtWidgets import QMessageBox, QProgressDialog
        from core.opus_downloader import (
            list_available_for_lang, OPUS_MODELS, is_downloaded,
            download_pairs_for_lang, estimate_ram_for_lang,
        )

        # 現在の対象言語
        tgt = getattr(self.interpreter, 'target_lang', 'en')
        pairs = list_available_for_lang(tgt)
        not_downloaded = [p for p in pairs if not is_downloaded(p)]

        if not not_downloaded:
            QMessageBox.information(self, "OPUS-MT",
                                    f"言語 [{tgt}] の全ペアはダウンロード済みです。")
            return

        ram_est = estimate_ram_for_lang(tgt)
        reply = QMessageBox.question(
            self, "OPUS-MT 一括ダウンロード",
            f"言語 [{tgt}] の関連ペア {len(not_downloaded)}個をダウンロードします。\n"
            f"（全ペアロード時 RAM ~{ram_est:.1f}GB）\n\n続行しますか？",
            QMessageBox.Yes | QMessageBox.No, QMessageBox.No)
        if reply != QMessageBox.Yes:
            return

        progress = QProgressDialog(f"OPUS-MT [{tgt}] 一括ダウンロード中...", "キャンセル", 0, 100, self)
        progress.setWindowTitle("一括ダウンロード")
        progress.setMinimumDuration(0)
        progress.setValue(0)

        import threading
        def _do_download():
            try:
                def on_progress(pair_key, ratio):
                    progress.setLabelText(f"OPUS-MT [{pair_key}] ダウンロード中...")
                    progress.setValue(int(ratio * 100))
                download_pairs_for_lang(tgt, on_progress=on_progress)
                progress.setValue(100)
            except Exception as e:
                print(f"[error] OPUS-MT一括ダウンロード失敗: {e}")
            finally:
                self._populate_opus_models()

        threading.Thread(target=_do_download, daemon=True).start()
        progress.exec()
        self.status_bar.showMessage(f"✅ OPUS-MT [{tgt}] 一括ダウンロード完了", 5000)

    def _setup_shortcuts(self):
        """全ショートカットをQShortcutで一元管理（メニュー非表示時も動作）"""
        for key, handler in [
            ("Ctrl+1", self._on_f1),               # ⌘1 ハイドモード
            ("Ctrl+2", self._on_f2),               # ⌘2 パニック
            ("Ctrl+3", self._toggle_embedded_panel), # ⌘3 画面切替
            ("Ctrl+5", self._toggle_stt),           # ⌘5 マイクON/OFF (メニュー用)
            ("Ctrl+6", lambda: self._set_stt_lang_mode("auto")),      # ⌘6 自動判定
            ("Ctrl+7", lambda: self._set_stt_lang_mode("attorney")),  # ⌘7 弁護人入力
            ("Ctrl+8", lambda: self._set_stt_lang_mode("defendant")), # ⌘8 相手入力
            ("Ctrl+D", self._on_simulate_defendant), # ⌘D テスト発言
            ("Ctrl+/", self._show_shortcut_help),   # ⌘/ ショートカット一覧
        ]:
            sc = QShortcut(QKeySequence(key), self)
            sc.setContext(Qt.WindowShortcut)
            sc.activated.connect(handler)

    def _setup_callbacks(self):
        def on_stream_token(utt: Utterance, token: str):
            self._interpreter_stream_token_ready.emit(utt, token)

        def on_utterance(utt: Utterance):
            self._interpreter_utterance_ready.emit(utt)

        self.interpreter.set_callbacks(
            on_utterance=on_utterance,
            on_stream_token=on_stream_token,
        )

    @Slot(object, str)
    def _on_interpreter_stream_token(self, utt: Utterance, token: str):
        self.stream_to_defendant.emit(token)
        if self._current_attorney_bubble:
            self._current_attorney_bubble.update_translation(utt.translated)

    @Slot(object)
    def _on_interpreter_utterance(self, utt: Utterance):
        if self._current_attorney_bubble:
            self._current_attorney_bubble.utterance = utt
            self._current_attorney_bubble.update_translation(utt.translated)
            self._current_attorney_bubble = None
        self.finish_defendant_stream.emit()
        self.status_bar.showMessage("待機中")

    def _start_translation_worker(self):
        worker = threading.Thread(target=self._translation_worker_loop, daemon=True)
        worker.start()
        self._translation_worker = worker

    def _next_translation_job(self, kind: str, text: str) -> TranslationJob:
        with self._translation_state_lock:
            self._translation_job_seq += 1
            return TranslationJob(
                job_id=self._translation_job_seq,
                session_token=self._translation_session_token,
                kind=kind,
                text=text,
            )

    def _enqueue_translation_job(self, kind: str, text: str) -> TranslationJob:
        job = self._next_translation_job(kind, text)
        self._translation_queue.put(job)
        return job

    def _is_translation_job_active(self, job: TranslationJob) -> bool:
        with self._translation_state_lock:
            return (
                job.session_token == self._translation_session_token
                and job.job_id not in self._cancelled_translation_job_ids
            )

    def _cancel_translation_job(self, job_id: int):
        with self._translation_state_lock:
            self._cancelled_translation_job_ids.add(job_id)

    def _invalidate_translation_jobs(self):
        with self._translation_state_lock:
            self._translation_session_token += 1
            self._cancelled_translation_job_ids.clear()

    def _translation_worker_loop(self):
        while True:
            job = self._translation_queue.get()
            try:
                if job is None:
                    return
                if not self._is_translation_job_active(job):
                    continue

                if job.kind in ("manual_attorney", "stt_attorney"):
                    utt = self._translate_attorney_text(job.text)
                    if not self._is_translation_job_active(job):
                        continue
                    if job.kind == "manual_attorney":
                        self._manual_attorney_translated.emit(job.job_id, utt)
                    else:
                        self._attorney_translated.emit(utt)
                elif job.kind == "defendant":
                    try:
                        utt = self.interpreter.translate_defendant(job.text)
                    except Exception as e:
                        import time as _t
                        utt = Utterance(
                            speaker=Speaker.DEFENDANT,
                            original=job.text,
                            translated=f"(翻訳エラー: {job.text})",
                            timestamp=_t.strftime("%H:%M"),
                        )
                        print(f"[DEF-TRANSLATE] ERROR: {e}")
                    if not self._is_translation_job_active(job):
                        continue
                    self._defendant_translated.emit(utt)
                else:
                    print(f"[warn] 不明な翻訳ジョブ種別: {job.kind}")
            except Exception as e:
                if not isinstance(job, TranslationJob) or not self._is_translation_job_active(job):
                    continue
                if job.kind == "manual_attorney":
                    self._manual_attorney_failed.emit(job.job_id, str(e))
                elif job.kind == "stt_attorney":
                    self._attorney_translation_failed.emit(job.text, str(e))
                elif job.kind == "defendant":
                    import time as _t
                    utt = Utterance(
                        speaker=Speaker.DEFENDANT,
                        original=job.text,
                        translated=f"(翻訳エラー: {job.text})",
                        timestamp=_t.strftime("%H:%M"),
                    )
                    self._defendant_translated.emit(utt)
            finally:
                if isinstance(job, TranslationJob):
                    with self._translation_state_lock:
                        self._cancelled_translation_job_ids.discard(job.job_id)
                self._translation_queue.task_done()

    # ----- 埋め込みパネル -----

    def _toggle_embedded_panel(self):
        """⌘3: 表示トグル"""
        if self._is_stealth:
            return  # ステルスモード中は無視
        if self._view_style == "switch":
            # ----- switch: 弁護人→被疑者→分割 を循環 -----
            self._switch_state = (self._switch_state + 1) % 3
            s = self._switch_state
            self._attorney_widget.setVisible(s != 1)   # 被疑者のみ時は非表示
            self._defendant_panel.setVisible(s != 0)    # 弁護人のみ時は非表示
            if s == 0:
                self._embed_action.setText("相手画面に切替  ⌘3")
                self.status_bar.showMessage("弁護人画面")
                self._embed_visible = False
            elif s == 1:
                self._embed_action.setText("左右分割  ⌘3")
                self.status_bar.showMessage("相手画面を表示中")
                self._embed_visible = True
            else:
                self._embed_action.setText("弁護人画面に戻る  ⌘3")
                self.status_bar.showMessage("左右分割表示")
                self._embed_visible = True
        else:
            # ----- split: QSplitter 表示/非表示 -----
            self._embed_visible = not self._embed_visible
            self._defendant_panel.setVisible(self._embed_visible)
            self._embed_action.setChecked(self._embed_visible)
            if self._embed_visible:
                current = self.width()
                if current < 1100:
                    self.resize(1200, self.height())
                self.status_bar.showMessage("相手画面: 埋め込み表示")
            else:
                self.status_bar.showMessage("相手画面: 非表示")

        self.embedded_panel_toggled.emit(self._embed_visible)

    def set_embedded_panel_visible(self, visible: bool):
        """外部から埋め込みパネルの表示を制御（main.pyから呼ぶ）"""
        self._embed_visible = visible
        if self._view_style == "switch":
            if visible:
                # 分割表示にする
                self._switch_state = 2
                self._attorney_widget.setVisible(True)
                self._defendant_panel.setVisible(True)
                self._embed_action.setText("弁護人画面に戻る  ⌘3")
            else:
                # 弁護人のみ
                self._switch_state = 0
                self._attorney_widget.setVisible(True)
                self._defendant_panel.setVisible(False)
                self._embed_action.setText("相手画面に切替  ⌘3")
        else:
            self._defendant_panel.setVisible(visible)
            self._embed_action.setChecked(visible)

    # ----- イベントハンドラ -----

    def _ensure_translation_available(self) -> bool:
        if self.interpreter.translation_ready:
            return True
        if self.interpreter.model_load_state == "loading":
            message = "⏳ 翻訳エンジン読込中 — まだ送信できません"
        elif self.interpreter.model_load_error:
            message = f"⚠ {self.interpreter.model_load_error}"
        else:
            message = "⚠ 翻訳エンジンが利用できません"
        self.status_bar.showMessage(message, 6000)
        return False

    def _translate_attorney_text(self, text: str) -> Utterance:
        import time as _t

        utt = Utterance(
            speaker=Speaker.ATTORNEY,
            original=text,
            timestamp=_t.strftime("%H:%M"),
        )
        tgt = self.interpreter.target_lang
        print(f"[ATT-TRANSLATE] start: text={text[:30]}, tgt={tgt}")

        processed_text, glossary_map = self.interpreter._glossary_pre_ja_to_foreign(text)
        if glossary_map:
            print(f"[ATT-TRANSLATE] glossary前処理: {processed_text}")

        if hasattr(self.interpreter.engine, "translate_detail"):
            from core.interpreter import detect_unknown_words
            result = self.interpreter.engine.translate_detail(processed_text, "ja", tgt)
            final_text = self.interpreter._glossary_post_ja_to_foreign(
                result.final_text, glossary_map
            )
            marked, unknowns = detect_unknown_words(text, final_text, "ja", tgt)
            utt.translated = marked
            utt.intermediate_en = result.intermediate_en
            utt.translation_route = result.route
            utt.unknown_words = unknowns
        else:
            raw = self.interpreter.engine.translate(processed_text, "ja", tgt)
            raw = self.interpreter._glossary_post_ja_to_foreign(raw, glossary_map)
            utt.translated = raw

        print(f"[ATT-TRANSLATE] done: translated={utt.translated[:50]}")
        return utt

    def _drop_pending_attorney_request(self, bubble: ConversationBubble):
        remove_key = None
        for request_id, pending_bubble in self._pending_attorney_bubbles.items():
            if pending_bubble is bubble:
                remove_key = request_id
                break
        if remove_key is not None:
            self._pending_attorney_bubbles.pop(remove_key, None)
            self._cancel_translation_job(remove_key)

    def _on_send_attorney(self):
        text = self.input_field.text().strip()
        if not text:
            return
        if not self._ensure_translation_available():
            return
        self.input_field.clear()
        self.input_field.setStyleSheet(self._input_normal_style)
        self.status_bar.showMessage("翻訳中...")

        utt = Utterance(speaker=Speaker.ATTORNEY, original=text, timestamp="--:--")
        import time as _t
        utt.timestamp = _t.strftime("%H:%M")
        utt.translated = "(翻訳中...)"
        bubble = ConversationBubble(utt, show_actions=True)
        bubble.edit_clicked.connect(self._on_bubble_edit)
        bubble.cancel_clicked.connect(self._on_bubble_cancel)
        self.log_layout.addWidget(bubble)
        self._last_attorney_bubble = bubble
        self._scroll_to_bottom()

        job = self._enqueue_translation_job("manual_attorney", text)
        self._pending_attorney_bubbles[job.job_id] = bubble

    # ------------------------------------------------------------------
    # 定型文テンプレート
    # ------------------------------------------------------------------

    def _load_phrases(self):
        """定型文を読み込む（.docx優先 → .json フォールバック）"""
        # 優先順: ユーザーdocx → ユーザーjson → 同梱json
        user_docx = os.path.expanduser("~/pli-models/定型文.docx")
        user_json = os.path.expanduser("~/pli-models/phrases.json")
        bundled_json = os.path.join(
            os.path.dirname(os.path.dirname(__file__)), "data", "phrases.json"
        )
        if os.path.exists(user_docx):
            try:
                return self._load_phrases_docx(user_docx)
            except Exception as e:
                print(f"[WARN] docx読込失敗, JSONにフォールバック: {e}")
        # JSON フォールバック
        import json
        path = user_json if os.path.exists(user_json) else bundled_json
        try:
            with open(path, "r", encoding="utf-8") as f:
                return json.load(f)
        except Exception:
            return {"categories": []}

    @staticmethod
    def _load_phrases_docx(path: str) -> dict:
        """Wordファイルから定型文を読み込む
        フォーマット:
          見出し1 = カテゴリ名（例: "⚖️ 権利告知"）
          表 = 2列（ラベル | 本文）
        """
        from docx import Document
        doc = Document(path)
        categories = []
        current_cat = None

        for element in doc.element.body:
            tag = element.tag.split("}")[-1]  # namespace除去

            if tag == "p":
                from docx.oxml.ns import qn
                pStyle = element.find(qn("w:pPr"))
                style_name = ""
                if pStyle is not None:
                    ps = pStyle.find(qn("w:pStyle"))
                    if ps is not None:
                        style_name = ps.get(qn("w:val"), "")

                if style_name in ("Heading1", "1", "heading 1", "見出し1", "見出し 1"):
                    # テキスト取得
                    text = element.text.strip() if hasattr(element, 'text') and element.text else ""
                    if not text:
                        # run から取得
                        texts = [r.text for r in element.findall(qn("w:r") + "/" + qn("w:t")) if r.text]
                        text = "".join(texts).strip()
                    if text:
                        # アイコンと名前を分離
                        icon, name = "", text
                        if len(text) > 1 and not text[0].isalnum() and not text[0] in "あ-ん":
                            # 先頭が絵文字っぽい
                            for i, ch in enumerate(text):
                                if ch == " " or ch == "　":
                                    icon = text[:i]
                                    name = text[i:].strip()
                                    break
                        current_cat = {"name": name, "icon": icon, "phrases": []}
                        categories.append(current_cat)

            elif tag == "tbl" and current_cat is not None:
                from docx.table import Table as DocxTable
                from docx.oxml.ns import qn
                rows = element.findall(qn("w:tr"))
                for row in rows:
                    cells = row.findall(qn("w:tc"))
                    if len(cells) >= 2:
                        label_texts = []
                        body_texts = []
                        for p in cells[0].findall(qn("w:p")):
                            for r in p.findall(qn("w:r")):
                                t = r.find(qn("w:t"))
                                if t is not None and t.text:
                                    label_texts.append(t.text)
                        for p in cells[1].findall(qn("w:p")):
                            for r in p.findall(qn("w:r")):
                                t = r.find(qn("w:t"))
                                if t is not None and t.text:
                                    body_texts.append(t.text)
                        label = "".join(label_texts).strip()
                        body = "".join(body_texts).strip()
                        # ヘッダー行スキップ
                        if label and body and label != "ラベル" and label != "名前":
                            current_cat["phrases"].append(
                                {"label": label, "text": body}
                            )
        return {"categories": categories}

    def _on_phrase_menu(self):
        """📋ボタン — 定型文カテゴリメニューをポップアップ"""
        data = self._load_phrases()
        menu = QMenu(self)
        menu.setStyleSheet(f"""
            QMenu {{
                background-color: {_SURFACE}; color: {_TEXT};
                font-size: 12px; border: 1px solid {_RAISED_D}; padding: 2px 0;
            }}
            QMenu::item {{ padding: 4px 20px 4px 10px; }}
            QMenu::item:selected {{ background-color: {_ACCENT}; color: white; }}
            QMenu::separator {{ height: 1px; background: {_RAISED_D}; margin: 2px 4px; }}
        """)
        for cat in data.get("categories", []):
            icon = cat.get("icon", "")
            sub = menu.addMenu(f'{icon} {cat["name"]}')
            sub.setStyleSheet(menu.styleSheet())
            for ph in cat.get("phrases", []):
                action = sub.addAction(ph["label"])
                action.triggered.connect(
                    lambda checked, t=ph["text"]: self._send_phrase(t)
                )
        menu.exec(self._phrase_btn.mapToGlobal(self._phrase_btn.rect().bottomLeft()))

    def _send_phrase(self, text: str):
        """定型文を翻訳して送信 — 入力欄経由で _on_send_attorney を再利用"""
        self.input_field.setText(text)
        self._on_send_attorney()

    def _save_phrases(self, data: dict):
        """定型文データを Word (.docx) で保存"""
        from docx import Document
        from docx.shared import Pt, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        user_path = os.path.expanduser("~/pli-models/定型文.docx")
        os.makedirs(os.path.dirname(user_path), exist_ok=True)

        doc = Document()

        # タイトル
        title = doc.add_heading("PLI 定型文テンプレート", level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        note = doc.add_paragraph(
            "このファイルを直接編集して保存すると、PLIの定型文メニューに反映されます。\n"
            "カテゴリ＝「見出し1」、フレーズ＝表（ラベル｜本文）の形式です。"
        )
        note.style = doc.styles["Normal"]
        for run in note.runs:
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

        for cat in data.get("categories", []):
            icon = cat.get("icon", "")
            name = cat.get("name", "")
            heading_text = f"{icon} {name}".strip() if icon else name
            doc.add_heading(heading_text, level=1)

            phrases = cat.get("phrases", [])
            if not phrases:
                doc.add_paragraph("（フレーズなし）")
                continue

            table = doc.add_table(rows=1 + len(phrases), cols=2)
            table.style = "Table Grid"

            # ヘッダー行
            hdr = table.rows[0]
            hdr.cells[0].text = "ラベル"
            hdr.cells[1].text = "本文"
            for cell in hdr.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
                        run.font.size = Pt(10)

            # データ行
            for i, ph in enumerate(phrases):
                row = table.rows[i + 1]
                row.cells[0].text = ph.get("label", "")
                row.cells[1].text = ph.get("text", "")
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(10)

            # 列幅調整（ラベル:本文 = 3:7）
            for row in table.rows:
                row.cells[0].width = Cm(4)
                row.cells[1].width = Cm(12)

            doc.add_paragraph("")  # 空行

        doc.save(user_path)
        self.status_bar.showMessage(f"定型文を保存しました: {user_path}")

    def _on_edit_phrases(self):
        """設定メニュー: 定型文編集ダイアログを開く"""
        data = self._load_phrases()
        dlg = PhraseEditorDialog(data, self._save_phrases, self)
        dlg.exec()

    def _on_edit_glossary(self):
        """設定メニュー: 固有名詞辞書編集ダイアログを開く"""
        dlg = GlossaryEditorDialog(self.interpreter, self)
        dlg.exec()

    # ------------------------------------------------------------------
    # 辞書検索
    # ------------------------------------------------------------------

    def _on_dict_dialog(self):
        """📖ボタン — 辞書検索ダイアログを開く"""
        dlg = DictionaryDialog(self.interpreter, self)
        dlg.exec()

    def _on_simulate_defendant(self):
        text, ok = QInputDialog.getText(
            self, "相手の発言（テスト）", "外国語で入力:",
        )
        if not ok or not text.strip():
            return
        self._process_defendant_speech(text.strip())

    def _process_defendant_speech(self, text: str):
        """被疑者発言を非同期翻訳（メインスレッドをブロックしない）"""
        if not self._ensure_translation_available():
            return
        self.status_bar.showMessage("翻訳中...")
        self.send_to_defendant.emit("defendant_echo", text)
        self._enqueue_translation_job("defendant", text)

    def _process_attorney_speech(self, text: str):
        """弁護人STT発言 → 即座に非同期翻訳 → 自動送信（修正/取消はバブル上で）"""
        if not self._ensure_translation_available():
            return
        self.status_bar.showMessage("🎤 文字起こし完了 — 翻訳中...")
        self._enqueue_translation_job("stt_attorney", text)

    def _on_attorney_translated(self, utt):
        """弁護人翻訳完了 → 即座に送信＋ログにバブル追加（修正/取消ボタン付き）"""
        utt.confirmed = True
        self.interpreter.conversation.append(utt)

        # 修正/取消ボタン付きバブルをログに追加
        bubble = ConversationBubble(utt, show_actions=True)
        bubble.edit_clicked.connect(self._on_bubble_edit)
        bubble.cancel_clicked.connect(self._on_bubble_cancel)
        self.log_layout.addWidget(bubble)
        self._last_attorney_bubble = bubble
        self._scroll_to_bottom()

        # 被疑者画面に即座に送信
        self.send_to_defendant.emit("attorney_start", utt.original)
        self.stream_to_defendant.emit(utt.translated)
        self.finish_defendant_stream.emit()
        self.status_bar.showMessage("送信済み — バブル上で修正/取消可能")

    @Slot(int, object)
    def _on_manual_attorney_translated(self, request_id: int, utt: Utterance):
        bubble = self._pending_attorney_bubbles.pop(request_id, None)
        if not bubble:
            return
        utt.confirmed = True
        self.interpreter.conversation.append(utt)
        bubble.utterance = utt
        bubble.update_translation(utt.translated)
        self._last_attorney_bubble = bubble
        self._scroll_to_bottom()

        self.send_to_defendant.emit("attorney_start", utt.original)
        self.stream_to_defendant.emit(utt.translated)
        self.finish_defendant_stream.emit()
        self.status_bar.showMessage("送信済み — バブル上で修正/取消可能")

    @Slot(int, str)
    def _on_manual_attorney_failed(self, request_id: int, message: str):
        bubble = self._pending_attorney_bubbles.pop(request_id, None)
        if not bubble:
            return
        bubble.utterance.translated = f"(翻訳エラー: {message})"
        bubble.update_translation(bubble.utterance.translated)
        self.status_bar.showMessage(f"翻訳エラー: {message}", 7000)

    @Slot(str, str)
    def _on_attorney_translation_failed(self, text: str, message: str):
        import time as _t

        utt = Utterance(
            speaker=Speaker.ATTORNEY,
            original=text,
            translated=f"(翻訳エラー: {message})",
            timestamp=_t.strftime("%H:%M"),
        )
        bubble = ConversationBubble(utt, show_actions=True)
        bubble.edit_clicked.connect(self._on_bubble_edit)
        bubble.cancel_clicked.connect(self._on_bubble_cancel)
        self.log_layout.addWidget(bubble)
        self._last_attorney_bubble = bubble
        self._scroll_to_bottom()
        self.status_bar.showMessage(f"翻訳エラー: {message}", 7000)

    def _on_bubble_edit(self, bubble):
        """バブル上の「修正」ボタン → 文字起こしを入力欄に戻して再翻訳可能にする"""
        utt = bubble.utterance
        self._drop_pending_attorney_request(bubble)
        # 会話履歴から除去
        if utt in self.interpreter.conversation:
            self.interpreter.conversation.remove(utt)
        # バブルをログから除去
        bubble.setParent(None)
        bubble.deleteLater()
        if self._last_attorney_bubble is bubble:
            self._last_attorney_bubble = None
        # 被疑者画面をクリア
        self.send_to_defendant.emit("clear_last", "")
        # 入力欄に文字起こしテキストを戻す
        self.input_field.setText(utt.original)
        self.input_field.setFocus()
        self.input_field.selectAll()
        self.status_bar.showMessage("文字起こしを修正してEnterで再送信")

    def _on_bubble_cancel(self, bubble):
        """バブル上の「取消」ボタン → バブルと会話履歴を除去"""
        utt = bubble.utterance
        self._drop_pending_attorney_request(bubble)
        # 会話履歴から除去
        if utt in self.interpreter.conversation:
            self.interpreter.conversation.remove(utt)
        # バブルをログから除去
        bubble.setParent(None)
        bubble.deleteLater()
        if self._last_attorney_bubble is bubble:
            self._last_attorney_bubble = None
        # 被疑者画面をクリア
        self.send_to_defendant.emit("clear_last", "")
        self.status_bar.showMessage("取り消しました")

    # ----- 被疑者バブル 修正/取消 -----

    def _on_def_bubble_edit(self, bubble):
        """被疑者バブルの「修正」→ 構文単語ごと修正ダイアログを開く"""
        utt = bubble.utterance
        self._editing_def_bubble = bubble
        english = utt.original
        self.interpreter.pause()
        self.defendant_correction.emit()
        chunks = self.interpreter.syntax_check(english)
        dialog = SyntaxCheckDialog(english, chunks, self.interpreter, self)
        dialog.confirmed.connect(self._on_def_syntax_confirmed)
        if dialog.exec() == QDialog.Rejected:
            self._editing_def_bubble = None
            self.interpreter.resume()

    def _on_def_syntax_confirmed(self, new_english: str, new_japanese: str):
        """被疑者バブルの構文修正確定 → バブル内の翻訳テキストを更新"""
        bubble = getattr(self, '_editing_def_bubble', None)
        if bubble and bubble.utterance:
            utt = bubble.utterance
            utt.original = new_english
            utt.translated = new_japanese
            bubble.update_translation(new_japanese)
            # 被疑者画面の最後のエントリも更新
            self.defendant_update_last.emit(new_english)
        self._editing_def_bubble = None
        self.interpreter.resume()
        self.status_bar.showMessage("相手の発言を修正しました")

    def _on_def_bubble_cancel(self, bubble):
        """被疑者バブルの「取消」→ バブルと会話履歴を除去"""
        utt = bubble.utterance
        if utt in self.interpreter.conversation:
            self.interpreter.conversation.remove(utt)
        bubble.setParent(None)
        bubble.deleteLater()
        if self._last_defendant_bubble is bubble:
            self._last_defendant_bubble = None
        # 被疑者画面からも除去
        self.send_to_defendant.emit("clear_last", "")
        self.status_bar.showMessage("相手の発言を取り消しました")

    def _on_defendant_translated(self, utt):
        """被疑者翻訳完了 → 即座にログに追加（修正/取消ボタン付き）"""
        print(f"[DEF-TRANSLATE] 弁護側ログに追加: {utt.original[:30]} → {utt.translated[:30]}")
        utt.confirmed = True
        self.interpreter.conversation.append(utt)

        # 修正/取消ボタン付きバブルをログに追加
        bubble = ConversationBubble(utt, show_actions=True)
        bubble.edit_clicked.connect(self._on_def_bubble_edit)
        bubble.cancel_clicked.connect(self._on_def_bubble_cancel)
        self.log_layout.addWidget(bubble)
        self._last_defendant_bubble = bubble
        self._scroll_to_bottom()

        # 被疑者画面に確認済み通知
        self.send_to_defendant.emit("defendant_confirmed", utt.original)
        self.status_bar.showMessage("相手の発言を追加 — バブル上で修正/取消可能")

    def _on_ok(self):
        if self._pending_utterance:
            utt = self._pending_utterance
            if self._pending_is_attorney:
                # 弁護人: ログに追加 + 被疑者画面に翻訳を送信
                utt.confirmed = True
                self.interpreter.conversation.append(utt)
                bubble = ConversationBubble(utt)
                self.log_layout.addWidget(bubble)
                self._scroll_to_bottom()
                # 被疑者画面に翻訳テキストを表示
                self.send_to_defendant.emit("attorney_start", utt.original)
                self.stream_to_defendant.emit(utt.translated)
                self.finish_defendant_stream.emit()
            else:
                # 被疑者: 従来通り
                self.interpreter.confirm_utterance(utt)
                bubble = ConversationBubble(utt)
                self.log_layout.addWidget(bubble)
                self._scroll_to_bottom()
                self.send_to_defendant.emit("defendant_confirmed", utt.original)
            self._pending_utterance = None
            self._pending_is_attorney = False
            self.approval_frame.setVisible(False)
            self.status_bar.showMessage("待機中")

    def _on_retry(self):
        was_attorney = self._pending_is_attorney
        self._pending_utterance = None
        self._pending_is_attorney = False
        self.approval_frame.setVisible(False)
        if not was_attorney:
            self.defendant_retry.emit()
        self.status_bar.showMessage("再発話待ち")

    def _on_manual_edit(self):
        if not self._pending_utterance:
            return
        if self._pending_is_attorney:
            # 弁護人: 文字起こしテキストを入力欄に入れて修正→再翻訳
            self.input_field.setText(self._pending_utterance.original)
            self.input_field.setFocus()
            self.input_field.selectAll()
            self._pending_utterance = None
            self._pending_is_attorney = False
            self.approval_frame.setVisible(False)
            self.status_bar.showMessage("文字起こしを修正してEnterで再送信")
            return
        self.interpreter.pause()
        self.defendant_correction.emit()
        english = self._pending_utterance.original
        chunks = self.interpreter.syntax_check(english)
        dialog = SyntaxCheckDialog(english, chunks, self.interpreter, self)
        dialog.confirmed.connect(self._on_syntax_confirmed)
        if dialog.exec() == QDialog.Rejected:
            self.interpreter.resume()
            return

    def _on_syntax_confirmed(self, new_english: str, new_japanese: str):
        if self._pending_utterance:
            self._pending_utterance.original = new_english
            self._pending_utterance.translated = new_japanese
            self._pending_utterance.confirmed = True
            bubble = ConversationBubble(self._pending_utterance)
            self.log_layout.addWidget(bubble)
            self._scroll_to_bottom()
            self.defendant_update_last.emit(new_english)
        self._pending_utterance = None
        self.approval_frame.setVisible(False)
        self.interpreter.resume()
        self.status_bar.showMessage("待機中")

    def _set_font_scale(self, scale: float):
        """フォントサイズスケール変更 → 全バブルを再構築"""
        _font_cfg.set_scale(scale)
        # 既存バブルを全て再構築
        self._rebuild_log_bubbles()
        # 被疑者パネルのフォントも連動
        if hasattr(self, '_defendant_panel') and self._defendant_panel:
            self._defendant_panel.refresh_fonts()
        self.status_bar.showMessage(f"フォントサイズ: ×{scale}", 3000)

    def _rebuild_log_bubbles(self):
        """会話ログのバブルをフォントサイズ変更後に再構築"""
        # 既存ウィジェットを全削除
        while self.log_layout.count():
            item = self.log_layout.takeAt(0)
            if item.widget():
                item.widget().deleteLater()
        self._last_attorney_bubble = None
        self._last_defendant_bubble = None
        # 会話履歴から再構築 — 全バブルに修正/取消ボタン
        for utt in self.interpreter.conversation:
            bubble = ConversationBubble(utt, show_actions=True)
            if utt.speaker == Speaker.ATTORNEY:
                bubble.edit_clicked.connect(self._on_bubble_edit)
                bubble.cancel_clicked.connect(self._on_bubble_cancel)
                self._last_attorney_bubble = bubble
            else:
                bubble.edit_clicked.connect(self._on_def_bubble_edit)
                bubble.cancel_clicked.connect(self._on_def_bubble_cancel)
                self._last_defendant_bubble = bubble
            self.log_layout.addWidget(bubble)
        self._scroll_to_bottom()

    def _on_language_change(self, lang_code: str):
        self.interpreter.set_target_language(lang_code)
        lang_name = get_language_name(lang_code)
        self.status_bar.showMessage(f"言語: {lang_name}")
        self.setWindowTitle(f"PLI — {lang_name}")
        self._lang_label.setText(lang_name)
        self.defendant_lang_change.emit(lang_code)

    def _on_rec_mode(self, mode: RecordMode):
        self.recorder.set_mode(mode)
        if mode == RecordMode.OFF:
            self.status_bar.showMessage("待機中")
            self.rec_size_label.setText("")
        elif mode == RecordMode.VOLATILE:
            self.recorder.start()
            self.status_bar.showMessage("REC (揮発)")
        elif mode == RecordMode.SAVE:
            self.recorder.start()
            self.status_bar.showMessage("REC (保存)")

    def _update_rec_size(self):
        if self.recorder.mode != RecordMode.OFF:
            size = self.recorder.get_buffer_size_mb()
            self.rec_size_label.setText(f"{size:.1f}MB")
        else:
            self.rec_size_label.setText("")

    def _on_f1(self):
        """⌘1: ハイドモード トグル"""
        self.request_hide.emit()  # main.pyのtoggle_hideが呼ばれる

    def _on_f2(self):
        self.request_panic.emit()

    def _on_save_json(self):
        """会話記録をJSON形式で保存"""
        if not self.interpreter.conversation:
            QMessageBox.information(self, "保存", "保存する会話がありません。")
            return
        import time as _t
        default_name = f"pli_record_{_t.strftime('%Y%m%d_%H%M%S')}.json"
        path, _ = QFileDialog.getSaveFileName(
            self, "会話記録を保存 (JSON)", default_name,
            "JSON Files (*.json);;All Files (*)"
        )
        if path:
            try:
                self.interpreter.save_conversation(path)
                self.status_bar.showMessage(f"保存完了: {path}")
            except Exception as e:
                QMessageBox.warning(self, "保存エラー", str(e))

    def _on_save_text(self):
        """会話記録をテキスト形式でエクスポート"""
        if not self.interpreter.conversation:
            QMessageBox.information(self, "エクスポート", "保存する会話がありません。")
            return
        import time as _t
        default_name = f"pli_record_{_t.strftime('%Y%m%d_%H%M%S')}.txt"
        path, _ = QFileDialog.getSaveFileName(
            self, "会話記録をエクスポート (テキスト)", default_name,
            "Text Files (*.txt);;All Files (*)"
        )
        if path:
            try:
                self.interpreter.export_conversation_text(path)
                self.status_bar.showMessage(f"エクスポート完了: {path}")
            except Exception as e:
                QMessageBox.warning(self, "エクスポートエラー", str(e))

    def _on_end_session(self):
        self.interpreter.clear_conversation()
        self.recorder.wipe()
        self.clear_logs()
        self._session_count += 1
        if self._session_label:
            self._session_label.setText(f"Session #{self._session_count:02d}")
        self.status_bar.showMessage("新セッション")

    def set_loading_state(self, loading: bool, ready: bool = True, message: str = ""):
        """モデル読み込み状態をUIに反映"""
        is_mock = self.interpreter.mock
        if loading:
            self._loading_got_real_progress = False
            self._loading_dots = 0
            if is_mock:
                self._loading_phase = "Whisper"
                self._loading_label.setText("⏳ Whisper読込中")
                self._mode_label.setText("MOCK")
            else:
                self._loading_phase = {
                    EngineType.LLM: "LLM",
                    EngineType.NLLB: "NLLB",
                    EngineType.HYBRID: "Hybrid",
                }.get(self.interpreter._engine_type, "LLM")
                self._loading_label.setText(f"⏳ {self._loading_phase}読込中")
                self._mode_label.setText("REAL")
            self.status_bar.showMessage("モデル読込中 — 音声認識は準備完了後に利用可能")
            # アニメーションタイマー（ドット循環）
            self._loading_anim_timer = QTimer()
            self._loading_anim_timer.timeout.connect(self._animate_loading)
            self._loading_anim_timer.start(600)
        else:
            # アニメーション停止
            if hasattr(self, '_loading_anim_timer') and self._loading_anim_timer:
                self._loading_anim_timer.stop()
                self._loading_anim_timer = None
            if ready:
                self._loading_label.setText("")
                if is_mock:
                    self._mode_label.setText("MOCK ✓")
                else:
                    self._mode_label.setText("REAL ✓")
                self.status_bar.showMessage("✓ モデル準備完了 — 翻訳・音声認識が利用可能", 5000)
            else:
                if self.interpreter.translation_ready and not self.interpreter.stt_ready:
                    self._loading_label.setText("⚠ STT不可")
                    self._mode_label.setText("MOCK !" if is_mock else "REAL !")
                else:
                    self._loading_label.setText("⚠ 読込失敗")
                    self._mode_label.setText("ERROR")
                self.status_bar.showMessage(
                    message or "モデルの読込に失敗しました", 8000
                )

    def _animate_loading(self):
        """読込中のドットアニメーション（実進捗が来るまで）"""
        if getattr(self, '_loading_got_real_progress', False):
            return  # 実進捗が来たらアニメ不要
        self._loading_dots = (self._loading_dots + 1) % 4
        dots = "." * self._loading_dots
        phase_name = getattr(self, '_loading_phase', 'LLM')
        self._loading_label.setText(f"⏳ {phase_name}読込中{dots}")

    def set_loading_progress(self, phase: str, progress: float):
        """モデル読み込みの進捗を更新（バックグラウンドスレッドからQTimerで呼ぶ）"""
        pct = int(progress * 100)
        phase_name = {
            "llm": "LLM",
            "nllb": "NLLB",
            "hybrid": "Hybrid",
            "stt": "Whisper",
        }.get(phase, phase.upper())
        self._loading_phase = phase_name
        if phase != "stt":
            if progress > 0.0:
                # 実進捗が来た — アニメーションから%表示に切替
                self._loading_got_real_progress = True
                self._loading_label.setText(f"⏳ {phase_name} {pct}%")
                self.status_bar.showMessage(f"{phase_name}モデル読込中… {pct}%")
            # progress == 0 はアニメーション続行（_animate_loadingに任せる）
        else:
            self._loading_got_real_progress = False  # Whisperはアニメに戻す
            if progress < 1.0:
                self._loading_label.setText("⏳ Whisper読込中…")
                self.status_bar.showMessage("音声認識モデル読込中…")
            else:
                self._loading_label.setText("✓ 準備完了")
                self.status_bar.showMessage("✓ 全モデル準備完了", 3000)

    def _toggle_stt(self):
        """⌘5: マイクON/OFF トグル"""
        # モデル未準備時はONにさせない
        if not self._stt_active and not self.interpreter._models_ready:
            if self.interpreter.model_load_state in ("error", "degraded") and self.interpreter.model_load_error:
                self.status_bar.showMessage(f"⚠ {self.interpreter.model_load_error}", 6000)
            else:
                self.status_bar.showMessage("⏳ モデル読込中 — 音声認識はまだ使えません", 3000)
            return

        self._stt_active = not self._stt_active
        self._stt_action.setChecked(self._stt_active)
        if self._stt_active:
            self._stt_action.setText("🎤 マイクOFF  Space")
            self.status_bar.showMessage("🎤 音声認識 ON — マイク待機中")
            self.input_field.setPlaceholderText("🎤 マイク待機中… / 日本語を入力してEnter")
            self._update_stt_mode_label()
        else:
            self._stt_action.setText("🎤 マイクON  Space")
            self.status_bar.showMessage("音声認識 OFF")
            self.input_field.setPlaceholderText("日本語を入力してEnter...")
            self.input_field.setStyleSheet(self._input_normal_style)
            self._stt_mode_label.setVisible(False)
        self.stt_toggled.emit(self._stt_active)

    def _set_stt_lang_mode(self, mode: str):
        """STT言語モード切替: auto / attorney / defendant"""
        self._stt_lang_mode = mode
        labels = {"auto": "自動判定", "attorney": "弁護人入力", "defendant": "相手入力"}
        # ラジオボタン同期
        for action in self._stt_lang_group.actions():
            if mode == "auto" and "AUTO" in action.text():
                action.setChecked(True)
            elif mode == "attorney" and "弁護人" in action.text():
                action.setChecked(True)
            elif mode == "defendant" and "相手" in action.text():
                action.setChecked(True)
        self.status_bar.showMessage(f"🎤 言語モード: {labels[mode]}", 3000)
        self._update_stt_mode_label()

    def _set_stt_sensitivity(self, preset: str):
        """マイク感度プリセット切替"""
        labels = {"high": "高感度", "normal": "標準", "low": "低感度"}
        self.status_bar.showMessage(f"🎚 マイク感度: {labels.get(preset, preset)}", 3000)
        self.stt_sensitivity_changed.emit(preset)

    def _set_stt_tempo(self, preset: str):
        """発話テンポプリセット切替"""
        labels = {"slow": "ゆっくり", "normal": "標準", "fast": "早口"}
        self.status_bar.showMessage(f"🗣 発話テンポ: {labels.get(preset, preset)}", 3000)
        self.stt_tempo_changed.emit(preset)

    def on_stt_result(self, text: str, lang: str):
        """STTリスナーからの結果を処理（main.pyから呼ばれる）"""
        if not text.strip():
            return

        # 言語モードによるオーバーライド
        mode = self._stt_lang_mode
        if mode == "attorney":
            is_attorney = True
        elif mode == "defendant":
            is_attorney = False
        else:
            # AUTO: Whisperの言語検出 vs ターゲット言語で判定
            # 弁護人=日本語、相手=ターゲット言語
            # Whisper検出が「ja」→ 弁護人
            # Whisper検出がターゲット言語 → 相手
            # どちらでもない → 日本語テキストか推定
            tgt = self.interpreter.target_lang
            if lang == "ja":
                is_attorney = True
            elif lang == tgt or (tgt == "en" and lang in ("en", "english")):
                is_attorney = False
            else:
                # 不明な場合: テキストに日本語文字が含まれるか判定
                import unicodedata
                ja_chars = sum(1 for c in text if unicodedata.name(c, "").startswith(("CJK", "HIRAGANA", "KATAKANA")))
                is_attorney = (ja_chars / max(len(text), 1)) > 0.3
            print(f"[STT-AUTO] lang={lang}, target={tgt}, is_attorney={is_attorney}, text={text[:30]}")

        if is_attorney:
            # 弁護人の発言 → 即座に翻訳→自動送信
            print(f"[STT] → attorney speech: {text[:40]}")
            self._process_attorney_speech(text)
        else:
            # 相手の発言 → 翻訳処理
            print(f"[STT] → defendant speech: {text[:40]}")
            self._process_defendant_speech(text)

    def on_stt_state_change(self, state_name: str):
        """STTリスナーの状態変更"""
        if not self._stt_active:
            return
        if state_name == "listening":
            self.status_bar.showMessage("🎤 発話検出中...")
            self.input_field.setPlaceholderText("🎤 発話検出中...")
        elif state_name == "processing":
            self.status_bar.showMessage("🎤 音声認識処理中...")
            self.input_field.setPlaceholderText("🎤 音声認識処理中...")
        elif state_name == "idle":
            self.status_bar.showMessage("🎤 音声認識 ON — マイク待機中")
            self.input_field.setPlaceholderText("🎤 マイク待機中… / 日本語を入力してEnter")

    def _update_stt_mode_label(self):
        """ステータスバーのSTTモード常設表示を更新"""
        if not self._stt_active:
            self._stt_mode_label.setVisible(False)
            return
        mode_text = {"auto": "🎤 AUTO", "attorney": "🎤 弁護人", "defendant": "🎤 相手"}
        self._stt_mode_label.setText(mode_text.get(self._stt_lang_mode, "🎤"))
        self._stt_mode_label.setVisible(True)
        # モードに応じて色分け
        if self._stt_lang_mode == "attorney":
            self._stt_mode_label.setStyleSheet(
                f"color: {_ACCENT}; font-size: 10px; font-weight: bold; padding: 0 6px;"
                f"border-left: 1px solid {_RAISED_D};"
            )
        elif self._stt_lang_mode == "defendant":
            self._stt_mode_label.setStyleSheet(
                f"color: {_DEF_CLR}; font-size: 10px; font-weight: bold; padding: 0 6px;"
                f"border-left: 1px solid {_RAISED_D};"
            )
        else:
            self._stt_mode_label.setStyleSheet(
                f"color: {_DIM}; font-size: 10px; padding: 0 6px;"
                f"border-left: 1px solid {_RAISED_D};"
            )

    def keyPressEvent(self, event):
        """キーボードショートカット（メニュー以外）"""
        focused = self.focusWidget()
        in_text_field = focused and (focused == self.input_field or
                                     isinstance(focused, QTextEdit))

        # 承認パネル表示中のショートカット (O/R/E)
        if self.approval_frame.isVisible() and not in_text_field:
            if event.key() == Qt.Key_O:
                self._on_ok()
                return
            elif event.key() == Qt.Key_R:
                self._on_retry()
                return
            elif event.key() == Qt.Key_E:
                self._on_manual_edit()
                return

        # スペースキーでマイクON/OFF（入力フィールドにフォーカスがないとき）
        if event.key() == Qt.Key_Space:
            if in_text_field:
                super().keyPressEvent(event)
                return
            self._toggle_stt()
            return

        # 入力欄にフォーカスがないとき、印字可能キーなら自動で入力欄にフォーカス移動
        if not in_text_field and event.text() and event.text().isprintable():
            self.input_field.setFocus()
            self.input_field.keyPressEvent(event)
            return

        super().keyPressEvent(event)

    def _scroll_to_bottom(self):
        def _do_scroll():
            sb = self.scroll_area.verticalScrollBar()
            sb.setValue(sb.maximum())
        # レイアウト更新後にスクロール（2段階で確実に）
        QTimer.singleShot(50, _do_scroll)
        QTimer.singleShot(200, _do_scroll)

    def _copy_all_conversation(self):
        """会話ログ全体をクリップボードにコピー"""
        from PySide6.QtWidgets import QApplication
        lines = []
        for i in range(self.log_layout.count()):
            widget = self.log_layout.itemAt(i).widget()
            if not isinstance(widget, ConversationBubble):
                continue
            u = widget.utterance
            tag = "弁護人" if u.speaker == Speaker.ATTORNEY else "相手"
            lines.append(f"[{u.timestamp}] {tag}: {u.original}")
            if u.intermediate_en:
                lines.append(f"  (EN) {u.intermediate_en}")
            if u.translated:
                lines.append(f"  → {u.translated}")
            lines.append("")
        if lines:
            QApplication.clipboard().setText("\n".join(lines))
            self.status_bar.showMessage("📋 全会話をコピーしました", 3000)
        else:
            self.status_bar.showMessage("コピーする会話がありません", 3000)

    def _show_about(self):
        """バージョン情報・開発者クレジット"""
        QMessageBox.about(
            self,
            "PLI について",
            "<div style='text-align:center;'>"
            "<h2 style='color:#1a3a6a;'>PLI — Private Link Interpreter</h2>"
            "<p style='font-size:13px;'>完全オフライン AI 通訳システム</p>"
            "<p style='font-size:13px;'>Version 2.0.0</p>"
            "<hr>"
            "<p style='font-size:12px; color:#666;'>開発</p>"
            "<p style='font-size:14px; font-weight:bold; color:#1a3a6a;'>"
            "中野通り法律事務所<br>弁護士  関  智之</p>"
            "<hr>"
            "<p style='font-size:10px; color:#999;'>"
            "Copyright &copy; 2025-2026 中野通り法律事務所 弁護士 関智之<br>"
            "All rights reserved.<br><br>"
            "法律用語辞書: 法務省 JLT v18.0 + DEA DIR-022-18</p>"
            "</div>"
        )

    def _show_shortcut_help(self):
        """ショートカット一覧ダイアログ"""
        dlg = QDialog(self)
        dlg.setWindowTitle("ショートカット一覧")
        dlg.setMinimumSize(420, 400)
        dlg.setStyleSheet(f"background-color: {_BG}; color: {_TEXT};")

        layout = QVBoxLayout(dlg)
        layout.setContentsMargins(16, 16, 16, 16)
        layout.setSpacing(8)

        title = QLabel("PLI キーボードショートカット")
        title.setStyleSheet(
            f"color: {_ACCENT}; font-size: 14px; font-weight: bold;"
            f"padding-bottom: 8px; border: none;"
        )
        layout.addWidget(title)

        shortcuts = [
            ("基本操作", [
                ("Enter", "入力した文章を送信"),
                ("Space", "マイクON/OFF（入力欄以外で）"),
            ]),
            ("画面制御", [
                ("⌘1", "画面を隠す（ステルスモード）"),
                ("⌘2", "緊急消去（パニック）"),
                ("⌘3", "相手画面の切替"),
            ]),
            ("音声認識", [
                ("⌘5", "マイクON/OFF"),
                ("⌘6", "言語モード: AUTO"),
                ("⌘7", "言語モード: 弁護人入力"),
                ("⌘8", "言語モード: 相手入力"),
                ("", "感度・テンポ → 音声認識メニューから"),
            ]),
            ("承認パネル表示中", [
                ("O", "OK（承認）"),
                ("R", "やり直し"),
                ("E", "手動修正"),
            ]),
            ("その他", [
                ("⌘D", "相手の発言をシミュレート"),
                ("⌘/", "このヘルプを表示"),
            ]),
        ]

        for section_title, items in shortcuts:
            section = QLabel(section_title)
            section.setStyleSheet(
                f"color: {_DIM}; font-size: 10px; letter-spacing: 2px;"
                f"padding-top: 6px; border: none;"
            )
            section.setFont(QFont("Menlo", 9))
            layout.addWidget(section)

            for key, desc in items:
                row = QHBoxLayout()
                row.setSpacing(12)
                key_label = QLabel(key)
                key_label.setFixedWidth(70)
                key_label.setAlignment(Qt.AlignRight | Qt.AlignVCenter)
                key_label.setStyleSheet(
                    f"color: {_TEXT}; font-size: 12px; font-weight: bold;"
                    f"font-family: Menlo; border: none;"
                    f"background-color: {_SURFACE}; padding: 2px 6px;"
                    f"border-top: 1px solid {_RAISED_L};"
                    f"border-left: 1px solid {_RAISED_L};"
                    f"border-bottom: 1px solid {_RAISED_D};"
                    f"border-right: 1px solid {_RAISED_D};"
                )
                desc_label = QLabel(desc)
                desc_label.setStyleSheet(
                    f"color: {_TEXT}; font-size: 12px; border: none;"
                )
                row.addWidget(key_label)
                row.addWidget(desc_label, stretch=1)
                layout.addLayout(row)

        layout.addStretch()
        close_btn = _make_btn("閉じる (Esc)", _BG)
        close_btn.clicked.connect(dlg.accept)
        layout.addWidget(close_btn, alignment=Qt.AlignRight)

        dlg.exec()

    def _create_stealth_widget(self) -> QWidget:
        """ハイドモード用ダミー画面 — メモ帳風"""
        w = QWidget()
        w.setStyleSheet("background-color: #ffffff;")
        layout = QVBoxLayout(w)
        layout.setContentsMargins(0, 0, 0, 0)

        # ダミーテキストエディタ（編集可能 — 実際にメモとして使える）
        editor = QTextEdit()
        editor.setStyleSheet("""
            QTextEdit {
                background-color: #ffffff;
                color: #333333;
                font-family: "Hiragino Mincho ProN", serif;
                font-size: 13px;
                border: none;
                padding: 20px;
            }
        """)
        editor.setPlainText(
            "令和6年度　第3回定例会議事録（案）\n\n"
            "日時：令和6年12月20日（金）14:00～15:30\n"
            "場所：第2会議室\n"
            "出席者：田中部長、山田課長、佐藤主任、鈴木\n\n"
            "1. 前回議事録の確認\n"
            "　前回の議事録について特に修正なく承認された。\n\n"
            "2. 年末年始の業務体制について\n"
            "　12月28日（土）から1月5日（日）まで休業とする。\n"
            "　緊急連絡先は田中部長の携帯電話とする。\n\n"
            "3. 来期の予算案について\n"
            "　各部門の概算要求を1月末までに取りまとめる。\n"
            "　詳細は別途配布の資料を参照のこと。\n\n"
            "4. その他\n"
            "　特になし。\n\n"
            "以上"
        )
        layout.addWidget(editor)
        return w

    def do_hide(self):
        """ハイドモード: ダミー画面に切替（ウィンドウは残す）"""
        self._is_stealth = True
        self._saved_title = self.windowTitle()

        # switchモード: 現在の表示状態を保存し、安全な状態にリセット
        if self._view_style == "switch":
            self._saved_switch_state = self._switch_state
            self._attorney_widget.setVisible(True)
            self._defendant_panel.setVisible(False)
            self._switch_state = 0

        self._stealth_stack.setCurrentIndex(1)
        self.setWindowTitle("メモ帳")
        self.menuBar().setVisible(False)
        self.status_bar.setVisible(False)

    def do_reveal(self):
        """ハイドモード解除: 実画面に復帰"""
        self._is_stealth = False
        self._stealth_stack.setCurrentIndex(0)

        # switchモード: 保存した表示状態を復元
        if self._view_style == "switch" and hasattr(self, '_saved_switch_state'):
            s = self._saved_switch_state
            self._switch_state = s
            self._attorney_widget.setVisible(s != 1)
            self._defendant_panel.setVisible(s != 0)

        self.setWindowTitle(getattr(self, '_saved_title', "PLI — Private Link Interpreter"))
        self.menuBar().setVisible(True)
        self.status_bar.setVisible(True)
        self.raise_()
        self.activateWindow()

    def clear_logs(self):
        self._invalidate_translation_jobs()
        self._pending_utterance = None
        self._pending_is_attorney = False
        self._current_attorney_bubble = None
        self._pending_attorney_bubbles.clear()
        self._last_attorney_bubble = None
        self._last_defendant_bubble = None
        self._editing_def_bubble = None
        self.approval_frame.setVisible(False)
        while self.log_layout.count():
            item = self.log_layout.takeAt(0)
            if item.widget():
                item.widget().deleteLater()
        self.clear_defendant.emit()

    def wipe_all(self, delete_saved_recordings: bool = False):
        self.interpreter.clear_conversation()
        self.recorder.wipe(delete_saved_files=delete_saved_recordings)
        self.clear_logs()
